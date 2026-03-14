from io import BytesIO

from docx import Document

from app.services.retrieval_eval_service import run_agentic_eval, run_retrieval_eval, run_v3_eval
from app.services.search_service import SearchService


def _create_project(client, *, name: str = "Retrieval Eval Project") -> dict:
    response = client.post(
        "/api/v1/projects",
        json={
            "name": name,
            "description": "Minimal retrieval evaluation fixtures",
            "default_external_policy": "allow_external",
        },
    )
    assert response.status_code == 201, response.text
    return response.json()["item"]


def _build_docx_bytes(*, paragraphs: list[str] | None = None, table_rows: list[list[str]] | None = None) -> bytes:
    document = Document()
    for paragraph in paragraphs or []:
        document.add_paragraph(paragraph)

    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for col_index, value in enumerate(row):
                table.cell(row_index, col_index).text = value

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_v21_minimal_retrieval_eval_set(client, monkeypatch):
    project = _create_project(client)
    docx_content = _build_docx_bytes(
        paragraphs=[
            "开题报告",
            "该系统面向室内空气质量检测与智能控制。",
            "报告还包含实施计划、创新点和优化建议。",
        ],
        table_rows=[
            ["题目", "基于STM32的室内空气质量检测与智能控制系统设计"],
            ["项目名称", "室内空气质量检测与智能控制系统"],
            ["创新点", "多传感器融合与自动控制联动"],
        ],
    )

    upload_response = client.post(
        f"/api/v1/projects/{project['id']}/sources/files",
        files=[
            (
                "files",
                (
                    "开题报告-刘艺.docx",
                    docx_content,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            )
        ],
    )
    assert upload_response.status_code == 201, upload_response.text

    monkeypatch.setattr("app.services.vector_store.VectorStore.search", lambda self, *, query, project_id, limit: [])

    service = SearchService()
    cases = [
        {
            "query": "我的题目是什么？",
            "history": None,
            "expect_grounded": True,
            "expect_second_pass": False,
        },
        {
            "query": "现在你知道了吗？",
            "history": [{"role": "user", "content_md": "我的开题报告题目是什么？"}],
            "expect_grounded": True,
            "expect_second_pass": True,
        },
        {
            "query": "你觉得报告有地方可以再优化吗",
            "history": None,
            "expect_grounded": True,
            "expect_second_pass": False,
        },
        {
            "query": "今天北京天气如何",
            "history": None,
            "expect_grounded": False,
            "expect_second_pass": True,
        },
    ]

    results = []
    for case in cases:
        evidence, diagnostics = service.retrieve_project_evidence_with_diagnostics(
            project["id"],
            case["query"],
            limit=3,
            apply_rerank=False,
            history=case["history"],
        )
        results.append(
            {
                "query": case["query"],
                "grounded_candidate": diagnostics["final"]["grounded_candidate"],
                "triggered_second_pass": diagnostics["triggered_second_pass"],
                "hit_count": len(evidence),
            }
        )

    assert results[0]["grounded_candidate"] is True
    assert results[0]["hit_count"] >= 1
    assert results[1]["grounded_candidate"] is True
    assert results[1]["triggered_second_pass"] is True
    assert results[2]["grounded_candidate"] is True
    assert results[3]["triggered_second_pass"] is True


def test_v24_retrieval_eval_runner_returns_structured_results(client, monkeypatch):
    monkeypatch.setattr("app.services.vector_store.VectorStore.search", lambda self, *, query, project_id, limit: [])

    report = run_retrieval_eval(client)

    assert report["case_count"] >= 6
    assert report["passed_case_count"] >= 5
    case_ids = {item["case_id"] for item in report["results"]}
    assert "direct_field_title" in case_ids
    assert "natural_follow_up" in case_ids
    assert "unrelated_weather" in case_ids

    natural_follow_up = next(item for item in report["results"] if item["case_id"] == "natural_follow_up")
    assert natural_follow_up["triggered_second_pass"] is True
    assert natural_follow_up["status"] == "passed"

    unrelated = next(item for item in report["results"] if item["case_id"] == "unrelated_weather")
    assert unrelated["grounded_candidate"] is False
    assert unrelated["status"] == "passed"


def test_v31_agentic_eval_runner_returns_v3_structured_results(client, monkeypatch):
    monkeypatch.setattr("app.services.vector_store.VectorStore.search", lambda self, *, query, project_id, limit: [])

    report = run_agentic_eval(client)

    assert report["case_count"] >= 7
    assert report["passed_case_count"] >= 6
    case_ids = {item["case_id"] for item in report["results"]}
    assert "project_only_grounding" in case_ids
    assert "project_plus_web_supplement" in case_ids
    assert "memory_assisted_follow_up" in case_ids
    assert "precheck_retry_project_once" in case_ids

    web_case = next(item for item in report["results"] if item["case_id"] == "project_plus_web_supplement")
    assert web_case["used_web"] is True
    assert web_case["primary_source_kind"] == "external_web"
    assert web_case["status"] == "passed"

    memory_case = next(item for item in report["results"] if item["case_id"] == "memory_assisted_follow_up")
    assert memory_case["memory_note_count"] >= 1
    assert memory_case["status"] == "passed"

    retry_case = next(item for item in report["results"] if item["case_id"] == "precheck_retry_project_once")
    assert retry_case["project_retry_count"] == 1
    assert len(retry_case["query_trace"]) >= 2
    assert retry_case["status"] == "passed"

    weak_case = next(item for item in report["results"] if item["case_id"] == "weak_source_mode_ready")
    assert weak_case["grounded_candidate"] is False
    assert weak_case["readiness_action"] == "insufficient"


def test_v31_combined_v3_eval_runner_includes_retrieval_and_agentic(client, monkeypatch):
    monkeypatch.setattr("app.services.vector_store.VectorStore.search", lambda self, *, query, project_id, limit: [])

    report = run_v3_eval(client)

    assert report["suite"] == "v3"
    assert report["retrieval"]["case_count"] >= 6
    assert report["agentic"]["case_count"] >= 7
    assert report["passed_case_count"] >= 11
