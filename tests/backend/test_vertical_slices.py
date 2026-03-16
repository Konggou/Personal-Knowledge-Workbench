from pathlib import Path
from io import BytesIO

from docx import Document
from app.services.search_service import SearchService
from app.api.routes.sessions import service as sessions_route_service


def test_healthcheck_reports_embedded_vector_backend(client):
    response = client.get("/api/v1/health")
    assert response.status_code == 200, response.text
    payload = response.json()

    assert payload["status"] == "ok"
    assert payload["qdrant_backend"]["backend_mode"] == "memory"
    assert payload["qdrant_backend"]["configured_qdrant_url"] == ":memory:"


def _create_project(client, *, name: str = "测试项目") -> dict:
    response = client.post(
        "/api/v1/projects",
        json={
            "name": name,
            "description": "后端聊天优先版回归",
            "default_external_policy": "allow_external",
        },
    )
    assert response.status_code == 201, response.text
    return response.json()["item"]


def _create_html_source(html_server: dict, filename: str, body: str) -> str:
    path = Path(html_server["root"]) / filename
    path.write_text(body, encoding="utf-8")
    return f"{html_server['base_url']}/{filename}"


def _build_docx_bytes(*, paragraphs: list[str] | None = None, table_rows: list[list[str]] | None = None) -> bytes:
    document = Document()
    for paragraph in paragraphs or []:
        document.add_paragraph(paragraph)

    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for col_index, value in enumerate(row):
                table.cell(row_index, col_index).text = value

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_structured_docx_bytes(
    *,
    headings: list[str] | None = None,
    paragraphs: list[str] | None = None,
    table_rows: list[list[str]] | None = None,
) -> bytes:
    document = Document()
    for heading in headings or []:
        document.add_heading(heading, level=1)
    for paragraph in paragraphs or []:
        document.add_paragraph(paragraph)

    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for col_index, value in enumerate(row):
                table.cell(row_index, col_index).text = value

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_project_session_chat_summary_and_report_flow(client, html_server):
    project = _create_project(client, name="会话主链项目")

    session_response = client.post(f"/api/v1/projects/{project['id']}/sessions")
    assert session_response.status_code == 201, session_response.text
    session = session_response.json()["item"]
    assert session["title"] == "新会话"

    source_url = _create_html_source(
        html_server,
        "chat.html",
        (
            "<html><head><title>Chat Source</title></head><body>"
            "<article><p>Lighthouse orchard benchmark is the anchor phrase.</p>"
            "<p>Roadmap comparison should always start from grounded evidence.</p></article>"
            "</body></html>"
        ),
    )
    source_response = client.post(
        f"/api/v1/projects/{project['id']}/sources/web",
        json={"url": source_url, "session_id": session["id"]},
    )
    assert source_response.status_code == 201, source_response.text

    message_response = client.post(
        f"/api/v1/sessions/{session['id']}/messages",
        json={"content": "What does the source say about lighthouse orchard benchmark?", "deep_research": False},
    )
    assert message_response.status_code == 200, message_response.text
    detail = message_response.json()["item"]

    assert detail["title"].startswith("What does the source say")
    answer = next(item for item in detail["messages"] if item["message_type"] == "assistant_answer")
    assert "lighthouse orchard benchmark" in answer["content_md"].lower()
    assert len(answer["sources"]) >= 1
    assert answer["source_mode"] == "project_grounded"

    summary_response = client.post(f"/api/v1/sessions/{session['id']}/summary")
    assert summary_response.status_code == 200, summary_response.text
    summary_detail = summary_response.json()["item"]
    summary_card = next(item for item in summary_detail["messages"] if item["message_type"] == "summary_card")
    assert "摘要" in summary_card["title"]

    report_response = client.post(f"/api/v1/sessions/{session['id']}/report")
    assert report_response.status_code == 200, report_response.text
    report_detail = report_response.json()["item"]
    report_card = next(item for item in report_detail["messages"] if item["message_type"] == "report_card")
    assert "# " in report_card["content_md"]

    delete_response = client.delete(f"/api/v1/messages/{report_card['id']}")
    assert delete_response.status_code == 200, delete_response.text
    after_delete = delete_response.json()["item"]
    assert all(item["id"] != report_card["id"] for item in after_delete["messages"])


def test_deep_research_stays_in_same_session_and_emits_events(client, html_server):
    project = _create_project(client, name="深度调研项目")
    session = client.post(f"/api/v1/projects/{project['id']}/sessions").json()["item"]

    source_url = _create_html_source(
        html_server,
        "research.html",
        (
            "<html><head><title>Research Source</title></head><body>"
            "<article><p>Tradeoff analysis needs evidence before conclusion.</p>"
            "<p>Compare roadmap options after retrieval, not before it.</p></article>"
            "</body></html>"
        ),
    )
    client.post(
        f"/api/v1/projects/{project['id']}/sources/web",
        json={"url": source_url, "session_id": session["id"]},
    )

    response = client.post(
        f"/api/v1/sessions/{session['id']}/messages",
        json={"content": "Compare the roadmap tradeoff and evaluate the options.", "deep_research": True},
    )
    assert response.status_code == 200, response.text
    detail = response.json()["item"]

    assert any(item["message_type"] == "status_card" and item["title"] == "正在查项目资料" for item in detail["messages"])
    assert any(item["message_type"] == "status_card" and item["title"] == "正在整理结论" for item in detail["messages"])
    assert any(item["message_type"] == "assistant_answer" and item["title"] == "调研结论" for item in detail["messages"])

    events_response = client.get(f"/api/v1/sessions/{session['id']}/events")
    assert events_response.status_code == 200, events_response.text
    assert "event: message" in events_response.text
    assert "event: status" in events_response.text
    assert "event: source_update" in events_response.text


def test_knowledge_grouping_and_source_lifecycle(client, html_server):
    project = _create_project(client, name="知识库项目")

    url_one = _create_html_source(
        html_server,
        "knowledge-one.html",
        "<html><head><title>First Knowledge</title></head><body><p>Alpha evidence line.</p></body></html>",
    )
    url_two = _create_html_source(
        html_server,
        "knowledge-two.html",
        "<html><head><title>Second Knowledge</title></head><body><p>Beta retrieval line.</p></body></html>",
    )
    first = client.post(f"/api/v1/projects/{project['id']}/sources/web", json={"url": url_one}).json()["item"]
    second = client.post(f"/api/v1/projects/{project['id']}/sources/web", json={"url": url_two}).json()["item"]

    knowledge_response = client.get("/api/v1/knowledge")
    assert knowledge_response.status_code == 200, knowledge_response.text
    groups = knowledge_response.json()["groups"]
    assert any(group["project_id"] == project["id"] for group in groups)

    search_response = client.get("/api/v1/knowledge?query=beta")
    assert search_response.status_code == 200, search_response.text
    filtered_groups = search_response.json()["groups"]
    assert any(item["title"] == "Second Knowledge" for group in filtered_groups for item in group["items"])

    archive_response = client.post(f"/api/v1/sources/{second['id']}/archive")
    assert archive_response.status_code == 200, archive_response.text
    archived = archive_response.json()["item"]
    assert archived["ingestion_status"] == "archived"

    restore_response = client.post(f"/api/v1/sources/{second['id']}/restore")
    assert restore_response.status_code == 200, restore_response.text
    restored = restore_response.json()["item"]
    assert restored["ingestion_status"] in {"ready", "ready_low_quality"}

    delete_response = client.delete(f"/api/v1/sources/{first['id']}")
    assert delete_response.status_code == 200, delete_response.text

    remaining_sources_response = client.get(f"/api/v1/projects/{project['id']}/sources")
    assert remaining_sources_response.status_code == 200, remaining_sources_response.text
    remaining_ids = [item["id"] for item in remaining_sources_response.json()["items"]]
    assert first["id"] not in remaining_ids


def test_deleted_archived_source_cannot_be_restored(client, html_server):
    project = _create_project(client, name="Lifecycle Guard Project")

    source_url = _create_html_source(
        html_server,
        "deleted-archived-source.html",
        "<html><head><title>Archived Source</title></head><body><p>Restore must not resurrect deleted sources.</p></body></html>",
    )
    source = client.post(f"/api/v1/projects/{project['id']}/sources/web", json={"url": source_url}).json()["item"]

    archive_response = client.post(f"/api/v1/sources/{source['id']}/archive")
    assert archive_response.status_code == 200, archive_response.text

    delete_response = client.delete(f"/api/v1/sources/{source['id']}")
    assert delete_response.status_code == 200, delete_response.text

    restore_response = client.post(f"/api/v1/sources/{source['id']}/restore")
    assert restore_response.status_code == 404, restore_response.text


def test_vector_search_still_returns_evidence_when_lexical_ranking_is_disabled(client, html_server, monkeypatch):
    monkeypatch.setattr(SearchService, "_score_chunks", lambda self, *, chunks, query, limit: [])

    project = _create_project(client, name="纯向量命中项目")

    source_url = _create_html_source(
        html_server,
        "vector.html",
        (
            "<html><head><title>Vector Source</title></head><body>"
            "<article><p>Lighthouse orchard benchmark is the main grounded retrieval phrase.</p>"
            "<p>Semantic retrieval should still find this line even if lexical ranking is disabled.</p></article>"
            "</body></html>"
        ),
    )
    client.post(f"/api/v1/projects/{project['id']}/sources/web", json={"url": source_url})

    results = SearchService().retrieve_project_evidence(
        project["id"],
        "What does the source say about lighthouse orchard benchmark?",
        limit=3,
        apply_rerank=False,
    )
    assert results
    assert any("lighthouse orchard benchmark" in item["normalized_text"].lower() for item in results)
    assert results[0]["source_title"] == "Vector Source"


def test_docx_table_content_is_searchable_for_chinese_sentence_queries(client):
    project = _create_project(client, name="DOCX Table Retrieval")

    docx_content = _build_docx_bytes(
        paragraphs=["河南科技大学毕业设计（论文）开题报告"],
        table_rows=[
            ["项目名称", "基于 Quest 3 的机械臂遥操作系统"],
            ["研究方向", "虚拟现实交互与遥操作"],
        ],
    )

    upload_response = client.post(
        f"/api/v1/projects/{project['id']}/sources/files",
        files=[
            (
                "files",
                (
                    "开题报告-刘艺.docx",
                    docx_content,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            )
        ],
    )
    assert upload_response.status_code == 201, upload_response.text

    results = SearchService().retrieve_project_evidence(project["id"], "我的开题报告是什么项目", limit=3, apply_rerank=False)

    assert results
    assert results[0]["source_title"] == "开题报告-刘艺.docx"
    assert "机械臂遥操作系统" in results[0]["normalized_text"]


def test_weak_source_mode_can_continue_with_llm_chat(client, monkeypatch):
    project = _create_project(client, name="弱资料模式项目")
    session = client.post(f"/api/v1/projects/{project['id']}/sessions").json()["item"]

    monkeypatch.setattr(sessions_route_service.llm, "is_configured", lambda: True)
    monkeypatch.setattr(
        sessions_route_service.llm,
        "generate_chat_reply",
        lambda *, conversation, research_mode=False, context_notes=None: "我是 Personal Knowledge Workbench 的通用对话助手，可以在没有资料命中时继续直接回答你。",
    )

    ask_response = client.post(
        f"/api/v1/sessions/{session['id']}/messages",
        json={"content": "你叫什么？", "deep_research": False},
    )
    assert ask_response.status_code == 200, ask_response.text
    detail = ask_response.json()["item"]

    answer = next(item for item in detail["messages"] if item["message_type"] == "assistant_answer")
    assert answer["source_mode"] == "weak_source_mode"
    assert answer["sources"] == []


def test_natural_chinese_title_question_can_fall_back_to_lexical_aliases(client, monkeypatch):
    project = _create_project(client, name="Natural Chinese Query")

    docx_content = _build_docx_bytes(
        paragraphs=["河南科技大学毕业设计（论文）开题报告"],
        table_rows=[
            ["课题名称", "基于STM32的室内空气质量检测与智能控制系统设计"],
            ["学生姓名", "刘艺"],
        ],
    )

    upload_response = client.post(
        f"/api/v1/projects/{project['id']}/sources/files",
        files=[
            (
                "files",
                (
                    "开题报告-刘艺.docx",
                    docx_content,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            )
        ],
    )
    assert upload_response.status_code == 201, upload_response.text

    monkeypatch.setattr("app.services.vector_store.VectorStore.search", lambda self, *, query, project_id, limit: [])

    results = SearchService().retrieve_project_evidence(
        project["id"],
        "现在你知道我的题目是什么了吗",
        limit=5,
        apply_rerank=False,
    )

    assert results
    assert results[0]["source_title"] == "开题报告-刘艺.docx"
    assert any("基于STM32的室内空气质量检测与智能控制系统设计" in item["normalized_text"] for item in results)

def test_docx_preview_exposes_structured_chunk_metadata(client):
    project = _create_project(client, name="Structured Chunk Preview")

    docx_content = _build_structured_docx_bytes(
        headings=["研究内容"],
        paragraphs=["系统需要覆盖空气质量采集、显示与报警。"],
        table_rows=[
            ["课题名称", "基于STM32的室内空气质量检测与智能控制系统设计"],
            ["预期成果", "完成硬件系统、控制程序与论文"],
        ],
    )

    upload_response = client.post(
        f"/api/v1/projects/{project['id']}/sources/files",
        files=[
            (
                "files",
                (
                    "structured-preview.docx",
                    docx_content,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            )
        ],
    )
    assert upload_response.status_code == 201, upload_response.text
    source = upload_response.json()["items"][0]

    preview_response = client.get(f"/api/v1/sources/{source['id']}")
    assert preview_response.status_code == 200, preview_response.text
    preview = preview_response.json()["item"]

    preview_chunks = preview["preview_chunks"]
    assert any(chunk["section_type"] == "heading" for chunk in preview_chunks)
    assert any(chunk["field_label"] == "课题名称" for chunk in preview_chunks)
    assert any(chunk["heading_path"] == "研究内容" for chunk in preview_chunks if chunk["section_type"] == "body")
    assert any(chunk["table_origin"] == "table_row_1" for chunk in preview_chunks if chunk["field_label"] == "课题名称")
    assert any(chunk["proposition_type"] == "method" for chunk in preview_chunks if chunk["section_type"] == "body")


def test_field_chunks_are_ranked_ahead_of_plain_body_hits_for_field_queries(client, monkeypatch):
    project = _create_project(client, name="Field Weighted Retrieval")

    docx_content = _build_structured_docx_bytes(
        headings=["研究内容"],
        paragraphs=[
            "这个系统围绕室内空气质量检测与智能控制展开。",
            "正文里也会提到课题名称和整体设计方向。",
        ],
        table_rows=[
            ["课题名称", "基于STM32的室内空气质量检测与智能控制系统设计"],
            ["研究内容", "完成采集、控制、显示与报警模块设计"],
        ],
    )

    upload_response = client.post(
        f"/api/v1/projects/{project['id']}/sources/files",
        files=[
            (
                "files",
                (
                    "field-weighted.docx",
                    docx_content,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            )
        ],
    )
    assert upload_response.status_code == 201, upload_response.text

    monkeypatch.setattr("app.services.vector_store.VectorStore.search", lambda self, *, query, project_id, limit: [])

    results = SearchService().retrieve_project_evidence(
        project["id"],
        "我的课题名称是什么",
        limit=3,
        apply_rerank=False,
    )

    assert results
    assert results[0]["section_type"] == "field"
    assert results[0]["field_label"] == "课题名称"
    assert "基于STM32的室内空气质量检测与智能控制系统设计" in results[0]["normalized_text"]


def test_project_delete_soft_deletes_and_excludes_from_list(client):
    """Test that project delete performs soft delete and excludes from default list."""
    # 1. Create a project
    project = _create_project(client, name="待删除项目")
    project_id = project["id"]

    # 2. Verify project exists and is active
    get_response = client.get(f"/api/v1/projects/{project_id}")
    assert get_response.status_code == 200, get_response.text
    assert get_response.json()["item"]["status"] == "active"

    # 3. Verify project appears in list
    list_response = client.get("/api/v1/projects")
    assert list_response.status_code == 200, list_response.text
    project_ids = [p["id"] for p in list_response.json()["items"]]
    assert project_id in project_ids

    # 4. Delete the project
    delete_response = client.delete(f"/api/v1/projects/{project_id}")
    assert delete_response.status_code == 200, delete_response.text
    deleted_project = delete_response.json()["item"]
    assert deleted_project["status"] == "archived"
    assert deleted_project["archived_at"] is not None

    # 5. Verify project no longer appears in default list (include_archived=false)
    list_after_delete = client.get("/api/v1/projects")
    assert list_after_delete.status_code == 200, list_after_delete.text
    project_ids_after = [p["id"] for p in list_after_delete.json()["items"]]
    assert project_id not in project_ids_after

    # 6. Verify project still exists when include_archived=true
    list_with_archived = client.get("/api/v1/projects?include_archived=true")
    assert list_with_archived.status_code == 200, list_with_archived.text
    archived_ids = [p["id"] for p in list_with_archived.json()["items"]]
    assert project_id in archived_ids

    # 7. Verify get project still returns the archived project
    get_after_delete = client.get(f"/api/v1/projects/{project_id}")
    assert get_after_delete.status_code == 200, get_after_delete.text
    assert get_after_delete.json()["item"]["status"] == "archived"


def test_project_archived_project_sessions_and_sources_hidden_from_lists(client, html_server):
    """Test that archived project's sessions and sources are hidden from all lists."""
    # 1. Create project with session and source
    project = _create_project(client, name="归档项目测试")
    project_id = project["id"]

    session = client.post(f"/api/v1/projects/{project_id}/sessions").json()["item"]
    session_id = session["id"]

    source_url = _create_html_source(
        html_server,
        "archive-test.html",
        "<html><head><title>Archive Test</title></head><body><p>Archive test content.</p></body></html>",
    )
    source = client.post(
        f"/api/v1/projects/{project_id}/sources/web",
        json={"url": source_url},
    ).json()["item"]
    source_id = source["id"]

    # 2. Verify session and source exist in project lists
    sessions_list = client.get(f"/api/v1/projects/{project_id}/sessions").json()["items"]
    assert any(s["id"] == session_id for s in sessions_list)

    sources_list = client.get(f"/api/v1/projects/{project_id}/sources").json()["items"]
    assert any(s["id"] == source_id for s in sources_list)

    # 3. Archive (delete) the project
    delete_response = client.delete(f"/api/v1/projects/{project_id}")
    assert delete_response.status_code == 200, delete_response.text

    # 4. Verify project is hidden from default list
    list_response = client.get("/api/v1/projects").json()["items"]
    assert not any(p["id"] == project_id for p in list_response)

    # 5. Verify sessions and sources are hidden from project lists
    sessions_after = client.get(f"/api/v1/projects/{project_id}/sessions").json()["items"]
    assert not any(s["id"] == session_id for s in sessions_after)

    sources_after = client.get(f"/api/v1/projects/{project_id}/sources").json()["items"]
    assert not any(s["id"] == source_id for s in sources_after)

    # 6. Verify sessions and sources are hidden from global lists
    global_sessions = client.get("/api/v1/sessions").json()["groups"]
    all_session_ids = []
    for group in global_sessions:
        all_session_ids.extend([s["id"] for s in group.get("sessions", [])])
    assert session_id not in all_session_ids

    global_sources_resp = client.get("/api/v1/knowledge").json()
    all_source_ids = []
    for group in global_sources_resp.get("groups", []):
        all_source_ids.extend([s["id"] for s in group.get("items", [])])
    assert source_id not in all_source_ids


def test_project_delete_returns_404_for_nonexistent_project(client):
    """Test that deleting a non-existent project returns 404."""
    response = client.delete("/api/v1/projects/non-existent-project-id")
    assert response.status_code == 404
    assert "not found" in response.json()["detail"].lower()
