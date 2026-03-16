from io import BytesIO

from docx import Document

from app.services.search_service import SearchService


def _build_docx_bytes(*, paragraphs: list[str] | None = None, table_rows: list[list[str]] | None = None) -> bytes:
    document = Document()
    for paragraph in paragraphs or []:
        document.add_paragraph(paragraph)

    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for col_index, value in enumerate(row):
                table.cell(row_index, col_index).text = value

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _create_project(client, *, name: str = "Retrieval V4 Project") -> dict:
    response = client.post(
        "/api/v1/projects",
        json={
            "name": name,
            "description": "Retrieval v4 tests",
            "default_external_policy": "allow_external",
        },
    )
    assert response.status_code == 201, response.text
    return response.json()["item"]


def test_rrf_promotes_shared_hits_above_single_channel_hits():
    service = SearchService()
    lexical_results = [
        {
            "chunk_id": "shared",
            "project_id": "project-1",
            "project_name": "Project",
            "source_id": "source-1",
            "source_title": "Shared",
            "source_type": "file_docx",
            "canonical_uri": "file:///shared.docx",
            "location_label": "Field #1",
            "excerpt": "shared lexical",
            "normalized_text": "shared lexical",
            "relevance_score": 1.4,
            "section_type": "field",
            "quality_level": "normal",
        },
        {
            "chunk_id": "lexical-only",
            "project_id": "project-1",
            "project_name": "Project",
            "source_id": "source-1",
            "source_title": "Lexical",
            "source_type": "file_docx",
            "canonical_uri": "file:///lexical.docx",
            "location_label": "Body #2",
            "excerpt": "lexical only",
            "normalized_text": "lexical only",
            "relevance_score": 1.2,
            "section_type": "body",
            "quality_level": "normal",
        },
    ]
    semantic_results = [
        {
            "chunk_id": "shared",
            "project_id": "project-1",
            "project_name": "Project",
            "source_id": "source-1",
            "source_title": "Shared",
            "source_type": "file_docx",
            "canonical_uri": "file:///shared.docx",
            "location_label": "Field #1",
            "excerpt": "shared semantic",
            "normalized_text": "shared semantic",
            "relevance_score": 0.81,
            "section_type": "field",
            "quality_level": "normal",
        },
        {
            "chunk_id": "semantic-only",
            "project_id": "project-1",
            "project_name": "Project",
            "source_id": "source-1",
            "source_title": "Semantic",
            "source_type": "file_docx",
            "canonical_uri": "file:///semantic.docx",
            "location_label": "Body #3",
            "excerpt": "semantic only",
            "normalized_text": "semantic only",
            "relevance_score": 0.79,
            "section_type": "body",
            "quality_level": "normal",
        },
    ]

    merged = service._merge_ranked_hits(semantic_results=semantic_results, lexical_results=lexical_results, limit=5)

    assert merged[0]["chunk_id"] == "shared"
    assert merged[0]["fusion_sources"] == ["lexical", "semantic"]


def test_rerank_falls_back_to_rule_when_cross_encoder_is_unavailable(monkeypatch):
    service = SearchService()
    monkeypatch.setattr(service.reranker, "_get_local_model", lambda: None)

    hits = [
        {
            "chunk_id": "chunk-1",
            "project_id": "project-1",
            "project_name": "Project",
            "source_id": "source-1",
            "source_title": "Indoor Air Quality Report",
            "source_type": "file_docx",
            "canonical_uri": "file:///report.docx",
            "location_label": "Body #1",
            "excerpt": "项目名称是室内空气质量检测系统。",
            "normalized_text": "项目名称是室内空气质量检测系统。",
            "relevance_score": 1.9,
            "section_type": "field",
            "quality_level": "normal",
        }
    ]

    reranked, diagnostics = service._rerank_hits(query="项目名称是什么", hits=hits, limit=3)

    assert reranked
    assert diagnostics["backend"] == "rule"
    assert diagnostics["applied"] is True


def test_invalid_fts_query_characters_do_not_break_retrieval(client, monkeypatch):
    project = _create_project(client)
    docx_content = _build_docx_bytes(
        paragraphs=["开题报告", "系统包含检测与控制两个模块。"],
        table_rows=[["题目", "室内空气质量检测系统设计"]],
    )
    upload_response = client.post(
        f"/api/v1/projects/{project['id']}/sources/files",
        files=[
            (
                "files",
                (
                    "retrieval-v4.docx",
                    docx_content,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            )
        ],
    )
    assert upload_response.status_code == 201, upload_response.text
    monkeypatch.setattr("app.services.vector_store.VectorStore.search", lambda self, *, query, project_id, limit: [])

    results = SearchService().retrieve_project_evidence(
        project["id"],
        '题目是什么 "(系统)" OR title?',
        limit=3,
        apply_rerank=False,
    )

    assert results
    assert results[0]["source_title"] == "retrieval-v4.docx"
