from __future__ import annotations

import json
import math
from dataclasses import dataclass
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
from time import perf_counter
from typing import Literal

from docx import Document

from app.core.settings import Settings, get_settings
from app.services.llm_service import LLMService
from app.services.reranker_service import RerankerService
from app.services.search_service import SearchService


RetrievalMode = Literal["lexical", "semantic", "hybrid"]
HydePolicy = Literal["off", "conditional", "all", "low_confidence"]


@dataclass(frozen=True)
class ChunkMatcher:
    source_title: str | None = None
    section_type: str | None = None
    field_label: str | None = None
    heading_path: str | None = None
    text_contains: str | None = None


@dataclass(frozen=True)
class BenchmarkCaseSpec:
    case_id: str
    query: str
    query_type: str
    history: list[dict] | None = None
    relevant_chunk_matchers: tuple[ChunkMatcher, ...] = ()
    relevant_source_titles: tuple[str, ...] = ()
    notes: str = ""


@dataclass(frozen=True)
class ResolvedBenchmarkCase:
    case_id: str
    query: str
    query_type: str
    history: list[dict] | None
    relevant_chunk_ids: tuple[str, ...]
    relevant_source_ids: tuple[str, ...]
    notes: str


@dataclass(frozen=True)
class RetrievalBenchmarkConfig:
    label: str
    stage: str
    retrieval_mode: RetrievalMode
    lexical_candidate_limit: int
    semantic_candidate_limit: int
    rrf_k: int
    rerank_enabled: bool
    reranker_top_n: int
    hyde_policy: HydePolicy
    final_retrieval_limit: int


class BenchmarkRerankerService(RerankerService):
    def __init__(self, settings: Settings) -> None:
        self._settings = settings

    @property
    def settings(self):
        return self._settings


class BenchmarkSearchService(SearchService):
    def __init__(self, *, config: RetrievalBenchmarkConfig, llm_service: LLMService | None = None) -> None:
        super().__init__(llm_service=llm_service)
        self.config = config
        self._benchmark_settings = _copy_settings(
            get_settings(),
            retrieval_lexical_candidate_limit=config.lexical_candidate_limit,
            retrieval_semantic_candidate_limit=config.semantic_candidate_limit,
            retrieval_rrf_k=config.rrf_k,
            reranker_top_n=config.reranker_top_n,
        )
        self.reranker = BenchmarkRerankerService(self._benchmark_settings)

    @property
    def settings(self):
        return self._benchmark_settings

    def _retrieve_ranked_hits(self, *, project_id: str, query: str, limit: int) -> list[dict]:
        available_chunks = self.repository.get_latest_chunks_for_project(project_id)
        lexical_results: list[dict] = []
        semantic_results: list[dict] = []

        if self.config.retrieval_mode in {"lexical", "hybrid"} and available_chunks:
            lexical_results = self._score_chunks(
                chunks=available_chunks,
                query=query,
                limit=max(limit, self.settings.retrieval_lexical_candidate_limit),
            )

        if self.config.retrieval_mode in {"semantic", "hybrid"}:
            semantic_results = self.vector_store.search(
                query=query,
                project_id=project_id,
                limit=max(limit, self.settings.retrieval_semantic_candidate_limit),
            )

        if self.config.retrieval_mode == "lexical":
            merged = lexical_results[:limit]
        elif self.config.retrieval_mode == "semantic":
            merged = semantic_results[:limit]
        else:
            merged = self._merge_ranked_hits(
                semantic_results=semantic_results,
                lexical_results=lexical_results,
                limit=limit,
            )

        return self._expand_structured_hits(
            query=query,
            hits=merged,
            available_chunks=available_chunks,
            limit=limit,
        )

    def _build_retry_queries(
        self,
        *,
        query: str,
        research_mode: bool,
        context_clues: list[str],
    ) -> list[tuple[str, str]]:
        retry_queries = super()._build_retry_queries(
            query=query,
            research_mode=research_mode,
            context_clues=context_clues,
        )

        if self.config.hyde_policy == "off":
            return [item for item in retry_queries if item[0] != "hyde_passage"]

        if self.config.hyde_policy in {"all", "low_confidence"}:
            filtered = [item for item in retry_queries if item[0] != "hyde_passage"]
            hypothetical_passage = self.llm.generate_hypothetical_passage(
                query=" ".join(query.split()).strip(),
                research_mode=research_mode,
            )
            if hypothetical_passage:
                filtered.append(("hyde_passage", hypothetical_passage))
            return filtered[: self.settings.retrieval_second_pass_limit]

        return retry_queries

    def _rerank_hits(self, *, query: str, hits: list[dict], limit: int) -> tuple[list[dict], dict]:
        if not self.config.rerank_enabled:
            return hits[:limit], {"backend": "disabled", "applied": False, "fallback_reason": "disabled_for_benchmark"}
        return super()._rerank_hits(query=query, hits=hits, limit=limit)


def run_offline_retrieval_benchmark(client, *, matrix: str = "full") -> dict:
    fixture = _seed_benchmark_fixture(client)
    cases = _build_resolved_cases(fixture)
    llm_service = LLMService()
    config_results: list[dict] = []

    stage_one_results = _evaluate_config_set(
        fixture=fixture,
        cases=cases,
        configs=_build_stage_one_configs(),
        llm_service=llm_service,
    )
    config_results.extend(stage_one_results)

    hybrid_results = _evaluate_config_set(
        fixture=fixture,
        cases=cases,
        configs=_build_hybrid_sweep_configs(matrix=matrix),
        llm_service=llm_service,
    )
    config_results.extend(hybrid_results)
    best_hybrid = _pick_best_config(hybrid_results)

    rerank_results = _evaluate_config_set(
        fixture=fixture,
        cases=cases,
        configs=_build_rerank_sweep_configs(best_hybrid=best_hybrid, matrix=matrix),
        llm_service=llm_service,
    )
    config_results.extend(rerank_results)
    best_rerank = _pick_best_config([item for item in rerank_results if item["config"]["rerank_enabled"]])
    rerank_baseline = next((item for item in rerank_results if item["config"]["label"] == "rerank-off-baseline"), None)

    hyde_results = _evaluate_config_set(
        fixture=fixture,
        cases=cases,
        configs=_build_hyde_sweep_configs(best_config=best_rerank, matrix=matrix),
        llm_service=llm_service,
    )
    config_results.extend(hyde_results)
    best_hyde = _pick_best_config(hyde_results)
    hyde_baseline = next((item for item in hyde_results if item["config"]["hyde_policy"] == "off"), None)

    summary = _build_summary(
        fixture=fixture,
        cases=cases,
        config_results=config_results,
        best_hybrid=best_hybrid,
        best_rerank=best_rerank,
        best_hyde=best_hyde,
        rerank_baseline=rerank_baseline,
        hyde_baseline=hyde_baseline,
    )
    artifacts_path = _write_artifacts(summary=summary, per_case=config_results)
    return {
        "summary": summary,
        "per_case": config_results,
        "artifacts_path": str(artifacts_path),
    }


def compute_case_metrics(
    *,
    retrieved_chunk_ids: list[str],
    retrieved_source_ids: list[str],
    relevant_chunk_ids: tuple[str, ...],
    relevant_source_ids: tuple[str, ...],
) -> dict:
    if relevant_chunk_ids:
        truth_ids = list(relevant_chunk_ids)
        predicted = retrieved_chunk_ids
        truth_type = "chunk"
    elif relevant_source_ids:
        truth_ids = list(relevant_source_ids)
        predicted = retrieved_source_ids
        truth_type = "source"
    else:
        return {
            "truth_type": "none",
            "recall_at": {str(k): None for k in (1, 3, 5, 10)},
            "hit_rate_at": {str(k): None for k in (1, 3, 5, 10)},
            "mrr_10": None,
            "ndcg_10": None,
            "top1_correct": None,
            "false_positive": bool(retrieved_chunk_ids or retrieved_source_ids),
        }

    truth_set = set(truth_ids)
    recall_at: dict[str, float] = {}
    hit_rate_at: dict[str, float] = {}
    for k in (1, 3, 5, 10):
        top_k = predicted[:k]
        hit_count = sum(1 for item in top_k if item in truth_set)
        recall_at[str(k)] = round(hit_count / max(len(truth_set), 1), 6)
        hit_rate_at[str(k)] = 1.0 if hit_count > 0 else 0.0

    reciprocal_rank = 0.0
    dcg = 0.0
    for index, item in enumerate(predicted[:10], start=1):
        if item in truth_set:
            if reciprocal_rank == 0.0:
                reciprocal_rank = 1.0 / index
            dcg += 1.0 / math.log2(index + 1)

    ideal_hits = min(len(truth_set), 10)
    idcg = sum(1.0 / math.log2(index + 1) for index in range(1, ideal_hits + 1))
    ndcg = (dcg / idcg) if idcg else 0.0

    return {
        "truth_type": truth_type,
        "recall_at": recall_at,
        "hit_rate_at": hit_rate_at,
        "mrr_10": round(reciprocal_rank, 6),
        "ndcg_10": round(ndcg, 6),
        "top1_correct": bool(predicted[:1] and predicted[0] in truth_set),
        "false_positive": False,
    }


def _evaluate_config_set(
    *,
    fixture: dict,
    cases: list[ResolvedBenchmarkCase],
    configs: list[RetrievalBenchmarkConfig],
    llm_service: LLMService,
) -> list[dict]:
    results: list[dict] = []
    for config in configs:
        service = BenchmarkSearchService(config=config, llm_service=llm_service)
        case_results: list[dict] = []
        for case in cases:
            started = perf_counter()
            evidence, diagnostics = service.retrieve_project_evidence_with_diagnostics(
                fixture["project"]["id"],
                case.query,
                limit=config.final_retrieval_limit,
                apply_rerank=config.rerank_enabled,
                history=case.history,
            )
            latency_ms = round((perf_counter() - started) * 1000, 3)
            retrieved_chunk_ids = [item["chunk_id"] for item in evidence if item.get("chunk_id")]
            retrieved_source_ids = _ordered_unique(item["source_id"] for item in evidence if item.get("source_id"))
            metrics = compute_case_metrics(
                retrieved_chunk_ids=retrieved_chunk_ids,
                retrieved_source_ids=retrieved_source_ids,
                relevant_chunk_ids=case.relevant_chunk_ids,
                relevant_source_ids=case.relevant_source_ids,
            )
            hyde_triggered = any(step.get("strategy") == "hyde_passage" for step in diagnostics.get("retry_steps", []))
            case_results.append(
                {
                    "case_id": case.case_id,
                    "query": case.query,
                    "query_type": case.query_type,
                    "history": case.history,
                    "relevant_chunk_ids": list(case.relevant_chunk_ids),
                    "relevant_source_ids": list(case.relevant_source_ids),
                    "retrieved_chunk_ids": retrieved_chunk_ids,
                    "retrieved_source_ids": retrieved_source_ids,
                    "retrieved_titles": [item.get("source_title") for item in evidence],
                    "hit_count": len(evidence),
                    "rerank_applied": bool(diagnostics.get("rerank", {}).get("applied")),
                    "hyde_triggered": hyde_triggered,
                    "latency_ms": latency_ms,
                    "metrics": metrics,
                    "diagnostics": diagnostics,
                    "notes": case.notes,
                }
            )

        results.append(
            {
                "config": {
                    "label": config.label,
                    "stage": config.stage,
                    "retrieval_mode": config.retrieval_mode,
                    "lexical_candidate_limit": config.lexical_candidate_limit,
                    "semantic_candidate_limit": config.semantic_candidate_limit,
                    "rrf_k": config.rrf_k,
                    "rerank_enabled": config.rerank_enabled,
                    "reranker_top_n": config.reranker_top_n,
                    "hyde_policy": config.hyde_policy,
                    "final_retrieval_limit": config.final_retrieval_limit,
                },
                "aggregate": _aggregate_case_results(case_results),
                "cases": case_results,
            }
        )
    return results


def _aggregate_case_results(case_results: list[dict]) -> dict:
    judged = [item for item in case_results if item["metrics"]["truth_type"] != "none"]
    unrelated = [item for item in case_results if item["metrics"]["truth_type"] == "none"]
    overall = {
        "recall_1": _mean(item["metrics"]["recall_at"]["1"] for item in judged),
        "recall_3": _mean(item["metrics"]["recall_at"]["3"] for item in judged),
        "recall_5": _mean(item["metrics"]["recall_at"]["5"] for item in judged),
        "recall_10": _mean(item["metrics"]["recall_at"]["10"] for item in judged),
        "hit_rate_1": _mean(item["metrics"]["hit_rate_at"]["1"] for item in judged),
        "hit_rate_3": _mean(item["metrics"]["hit_rate_at"]["3"] for item in judged),
        "hit_rate_5": _mean(item["metrics"]["hit_rate_at"]["5"] for item in judged),
        "hit_rate_10": _mean(item["metrics"]["hit_rate_at"]["10"] for item in judged),
        "mrr_10": _mean(item["metrics"]["mrr_10"] for item in judged),
        "ndcg_10": _mean(item["metrics"]["ndcg_10"] for item in judged),
        "top1_accuracy": _mean(1.0 if item["metrics"]["top1_correct"] else 0.0 for item in judged),
        "false_positive_rate": _mean(1.0 if item["metrics"]["false_positive"] else 0.0 for item in unrelated),
        "hyde_trigger_rate": _mean(1.0 if item["hyde_triggered"] else 0.0 for item in case_results),
        "hyde_trigger_precision": _mean(
            item["metrics"]["hit_rate_at"]["5"]
            for item in case_results
            if item["hyde_triggered"] and item["metrics"]["truth_type"] != "none"
        ),
        "latency_avg_ms": _mean(item["latency_ms"] for item in case_results),
        "latency_p95_ms": _percentile([item["latency_ms"] for item in case_results], 0.95),
        "score": None,
    }
    overall["score"] = round(
        (
            (overall["recall_5"] or 0.0) * 0.35
            + (overall["mrr_10"] or 0.0) * 0.25
            + (overall["ndcg_10"] or 0.0) * 0.25
            + (overall["top1_accuracy"] or 0.0) * 0.15
            - (overall["false_positive_rate"] or 0.0) * 0.2
        ),
        6,
    )
    by_query_type: dict[str, dict] = {}
    for query_type in sorted({item["query_type"] for item in case_results}):
        subset = [item for item in case_results if item["query_type"] == query_type]
        by_query_type[query_type] = {
            "case_count": len(subset),
            "recall_5": _mean(
                item["metrics"]["recall_at"]["5"] for item in subset if item["metrics"]["truth_type"] != "none"
            ),
            "mrr_10": _mean(
                item["metrics"]["mrr_10"] for item in subset if item["metrics"]["truth_type"] != "none"
            ),
            "false_positive_rate": _mean(
                1.0 if item["metrics"]["false_positive"] else 0.0
                for item in subset
                if item["metrics"]["truth_type"] == "none"
            ),
        }
    return {"overall": overall, "by_query_type": by_query_type}


def _build_summary(
    *,
    fixture: dict,
    cases: list[ResolvedBenchmarkCase],
    config_results: list[dict],
    best_hybrid: dict,
    best_rerank: dict,
    best_hyde: dict,
    rerank_baseline: dict | None,
    hyde_baseline: dict | None,
) -> dict:
    settings = get_settings()
    default_supported = {
        "lexical_candidate_limit": best_hybrid["config"]["lexical_candidate_limit"] == settings.retrieval_lexical_candidate_limit,
        "semantic_candidate_limit": best_hybrid["config"]["semantic_candidate_limit"] == settings.retrieval_semantic_candidate_limit,
        "rrf_k": best_hybrid["config"]["rrf_k"] == settings.retrieval_rrf_k,
        "reranker_top_n": best_rerank["config"]["reranker_top_n"] == settings.reranker_top_n,
        "final_retrieval_limit": best_hyde["config"]["final_retrieval_limit"] == 3,
    }
    recommendations = {
        "retrieval_mode": best_hybrid["config"]["retrieval_mode"],
        "lexical_candidate_limit": best_hybrid["config"]["lexical_candidate_limit"],
        "semantic_candidate_limit": best_hybrid["config"]["semantic_candidate_limit"],
        "rrf_k": best_hybrid["config"]["rrf_k"],
        "reranker_top_n": best_rerank["config"]["reranker_top_n"],
        "hyde_policy": best_hyde["config"]["hyde_policy"],
        "final_retrieval_limit": best_hyde["config"]["final_retrieval_limit"],
    }
    rerank_delta = _compare_aggregates(rerank_baseline, best_rerank)
    hyde_delta = _compare_aggregates(hyde_baseline, best_hyde)
    report_markdown = _build_report_markdown(
        recommendations=recommendations,
        best_hybrid=best_hybrid,
        best_rerank=best_rerank,
        best_hyde=best_hyde,
        rerank_delta=rerank_delta,
        hyde_delta=hyde_delta,
        default_supported=default_supported,
    )
    return {
        "suite": "offline_retrieval_benchmark",
        "generated_at": datetime.now(UTC).isoformat(),
        "project_id": fixture["project"]["id"],
        "case_count": len(cases),
        "query_type_counts": _count_by_query_type(cases),
        "config_count": len(config_results),
        "recommendations": recommendations,
        "default_parameter_support": default_supported,
        "best_configs": {
            "hybrid": {"config": best_hybrid["config"], "aggregate": best_hybrid["aggregate"]},
            "rerank": {"config": best_rerank["config"], "aggregate": best_rerank["aggregate"]},
            "hyde": {"config": best_hyde["config"], "aggregate": best_hyde["aggregate"]},
        },
        "deltas": {
            "rerank_vs_off": rerank_delta,
            "hyde_vs_off": hyde_delta,
        },
        "configs": [{"config": item["config"], "aggregate": item["aggregate"]} for item in config_results],
        "report_markdown": report_markdown,
    }


def _build_report_markdown(
    *,
    recommendations: dict,
    best_hybrid: dict,
    best_rerank: dict,
    best_hyde: dict,
    rerank_delta: dict | None,
    hyde_delta: dict | None,
    default_supported: dict,
) -> str:
    lines = [
        "# Offline Retrieval Benchmark",
        "",
        "## Recommended Parameters",
        f"- Retrieval mode: `{recommendations['retrieval_mode']}`",
        f"- Lexical candidate limit: `{recommendations['lexical_candidate_limit']}`",
        f"- Semantic candidate limit: `{recommendations['semantic_candidate_limit']}`",
        f"- RRF k: `{recommendations['rrf_k']}`",
        f"- Reranker top_n: `{recommendations['reranker_top_n']}`",
        f"- HyDE policy: `{recommendations['hyde_policy']}`",
        f"- Final retrieval limit: `{recommendations['final_retrieval_limit']}`",
        "",
        "## Why These Values",
        f"- Best hybrid score: `{best_hybrid['aggregate']['overall']['score']}`",
        f"- Best rerank score: `{best_rerank['aggregate']['overall']['score']}`",
        f"- Best HyDE score: `{best_hyde['aggregate']['overall']['score']}`",
    ]
    if rerank_delta:
        lines.append(f"- Rerank delta recall@5: `{rerank_delta['recall_5_delta']}`")
        lines.append(f"- Rerank delta MRR@10: `{rerank_delta['mrr_10_delta']}`")
        lines.append(f"- Rerank delta top1 accuracy: `{rerank_delta['top1_accuracy_delta']}`")
    if hyde_delta:
        lines.append(f"- HyDE delta recall@5: `{hyde_delta['recall_5_delta']}`")
        lines.append(f"- HyDE delta false positive rate: `{hyde_delta['false_positive_rate_delta']}`")
    lines.extend(
        [
            "",
            "## Default Parameter Support",
            f"- lexical candidate limit: `{default_supported['lexical_candidate_limit']}`",
            f"- semantic candidate limit: `{default_supported['semantic_candidate_limit']}`",
            f"- rrf_k: `{default_supported['rrf_k']}`",
            f"- reranker_top_n: `{default_supported['reranker_top_n']}`",
            f"- final retrieval limit=3: `{default_supported['final_retrieval_limit']}`",
        ]
    )
    return "\n".join(lines)


def _write_artifacts(*, summary: dict, per_case: list[dict]) -> Path:
    settings = get_settings()
    artifact_dir = settings.data_dir / "evals" / "retrieval-benchmark" / datetime.now(UTC).strftime("%Y%m%d-%H%M%S")
    artifact_dir.mkdir(parents=True, exist_ok=True)
    (artifact_dir / "summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    (artifact_dir / "per_case.json").write_text(json.dumps(per_case, ensure_ascii=False, indent=2), encoding="utf-8")
    (artifact_dir / "report.md").write_text(summary["report_markdown"], encoding="utf-8")
    return artifact_dir


def _build_stage_one_configs() -> list[RetrievalBenchmarkConfig]:
    settings = get_settings()
    return [
        RetrievalBenchmarkConfig(
            label="lexical-only",
            stage="retrieval_modes",
            retrieval_mode="lexical",
            lexical_candidate_limit=settings.retrieval_lexical_candidate_limit,
            semantic_candidate_limit=settings.retrieval_semantic_candidate_limit,
            rrf_k=settings.retrieval_rrf_k,
            rerank_enabled=False,
            reranker_top_n=settings.reranker_top_n,
            hyde_policy="off",
            final_retrieval_limit=3,
        ),
        RetrievalBenchmarkConfig(
            label="semantic-only",
            stage="retrieval_modes",
            retrieval_mode="semantic",
            lexical_candidate_limit=settings.retrieval_lexical_candidate_limit,
            semantic_candidate_limit=settings.retrieval_semantic_candidate_limit,
            rrf_k=settings.retrieval_rrf_k,
            rerank_enabled=False,
            reranker_top_n=settings.reranker_top_n,
            hyde_policy="off",
            final_retrieval_limit=3,
        ),
        RetrievalBenchmarkConfig(
            label="hybrid-current",
            stage="retrieval_modes",
            retrieval_mode="hybrid",
            lexical_candidate_limit=settings.retrieval_lexical_candidate_limit,
            semantic_candidate_limit=settings.retrieval_semantic_candidate_limit,
            rrf_k=settings.retrieval_rrf_k,
            rerank_enabled=False,
            reranker_top_n=settings.reranker_top_n,
            hyde_policy="off",
            final_retrieval_limit=3,
        ),
    ]


def _build_hybrid_sweep_configs(*, matrix: str) -> list[RetrievalBenchmarkConfig]:
    values = {
        "full": {"candidate_limits": (8, 16, 24, 40), "rrf_k_values": (10, 30, 50, 80)},
        "smoke": {"candidate_limits": (8, 24), "rrf_k_values": (10, 50)},
    }[matrix]
    configs: list[RetrievalBenchmarkConfig] = []
    for lexical_limit in values["candidate_limits"]:
        for semantic_limit in values["candidate_limits"]:
            for rrf_k in values["rrf_k_values"]:
                configs.append(
                    RetrievalBenchmarkConfig(
                        label=f"hybrid-l{lexical_limit}-s{semantic_limit}-k{rrf_k}",
                        stage="hybrid_sweep",
                        retrieval_mode="hybrid",
                        lexical_candidate_limit=lexical_limit,
                        semantic_candidate_limit=semantic_limit,
                        rrf_k=rrf_k,
                        rerank_enabled=False,
                        reranker_top_n=get_settings().reranker_top_n,
                        hyde_policy="off",
                        final_retrieval_limit=3,
                    )
                )
    return configs


def _build_rerank_sweep_configs(*, best_hybrid: dict, matrix: str) -> list[RetrievalBenchmarkConfig]:
    top_n_values = (4, 8, 12, 16) if matrix == "full" else (4, 8)
    config = best_hybrid["config"]
    configs = [
        RetrievalBenchmarkConfig(
            label="rerank-off-baseline",
            stage="rerank_sweep",
            retrieval_mode=config["retrieval_mode"],
            lexical_candidate_limit=config["lexical_candidate_limit"],
            semantic_candidate_limit=config["semantic_candidate_limit"],
            rrf_k=config["rrf_k"],
            rerank_enabled=False,
            reranker_top_n=get_settings().reranker_top_n,
            hyde_policy="off",
            final_retrieval_limit=config["final_retrieval_limit"],
        )
    ]
    for top_n in top_n_values:
        configs.append(
            RetrievalBenchmarkConfig(
                label=f"rerank-topn-{top_n}",
                stage="rerank_sweep",
                retrieval_mode=config["retrieval_mode"],
                lexical_candidate_limit=config["lexical_candidate_limit"],
                semantic_candidate_limit=config["semantic_candidate_limit"],
                rrf_k=config["rrf_k"],
                rerank_enabled=True,
                reranker_top_n=top_n,
                hyde_policy="off",
                final_retrieval_limit=config["final_retrieval_limit"],
            )
        )
    return configs


def _build_hyde_sweep_configs(*, best_config: dict, matrix: str) -> list[RetrievalBenchmarkConfig]:
    policies: tuple[HydePolicy, ...] = ("off", "conditional", "all", "low_confidence")
    if matrix == "smoke":
        policies = ("off", "conditional")
    config = best_config["config"]
    return [
        RetrievalBenchmarkConfig(
            label=f"hyde-{policy}",
            stage="hyde_sweep",
            retrieval_mode=config["retrieval_mode"],
            lexical_candidate_limit=config["lexical_candidate_limit"],
            semantic_candidate_limit=config["semantic_candidate_limit"],
            rrf_k=config["rrf_k"],
            rerank_enabled=config["rerank_enabled"],
            reranker_top_n=config["reranker_top_n"],
            hyde_policy=policy,
            final_retrieval_limit=config["final_retrieval_limit"],
        )
        for policy in policies
    ]


def _pick_best_config(results: list[dict]) -> dict:
    return max(results, key=lambda item: item["aggregate"]["overall"]["score"])


def _compare_aggregates(baseline: dict | None, candidate: dict | None) -> dict | None:
    if baseline is None or candidate is None:
        return None
    base = baseline["aggregate"]["overall"]
    current = candidate["aggregate"]["overall"]
    return {
        "recall_5_delta": round((current["recall_5"] or 0.0) - (base["recall_5"] or 0.0), 6),
        "mrr_10_delta": round((current["mrr_10"] or 0.0) - (base["mrr_10"] or 0.0), 6),
        "top1_accuracy_delta": round((current["top1_accuracy"] or 0.0) - (base["top1_accuracy"] or 0.0), 6),
        "false_positive_rate_delta": round((current["false_positive_rate"] or 0.0) - (base["false_positive_rate"] or 0.0), 6),
    }


def _seed_benchmark_fixture(client) -> dict:
    project_response = client.post(
        "/api/v1/projects",
        json={
            "name": "Offline Retrieval Benchmark Project",
            "description": "Offline retrieval benchmark fixtures",
            "default_external_policy": "allow_external",
        },
    )
    if project_response.status_code != 201:
        raise RuntimeError(f"Failed to create benchmark project: {project_response.text}")
    project = project_response.json()["item"]

    documents = [
        (
            "benchmark-overview.docx",
            _build_docx(
                paragraphs=[
                    "系统面向室内空气质量检测与智能控制。",
                    "研究内容包括采集模块、控制模块、显示模块与报警模块。",
                    "控制模块负责根据空气质量指标联动风扇转速。",
                ],
                table_rows=[
                    ["题目", "基于STM32的室内空气质量检测与智能控制系统设计"],
                    ["项目名称", "室内空气质量检测与智能控制系统"],
                    ["创新点", "多传感器融合与自动控制联动"],
                ],
            ),
        ),
        (
            "benchmark-optimization.docx",
            _build_docx(
                paragraphs=[
                    "优化建议章节建议继续优化多传感器融合。",
                    "还应补充实验验证并降低误报率。",
                    "报告中可增加异常场景测试与长期稳定性测试。",
                ],
                table_rows=[
                    ["优化建议", "继续优化多传感器融合、补充实验验证、降低误报率"],
                ],
            ),
        ),
        (
            "benchmark-feasibility.docx",
            _build_docx(
                paragraphs=[
                    "结论章节指出该方案具备实现可行性。",
                    "但仍需要补充实验验证与行业 benchmark 对照。",
                ],
                table_rows=[
                    ["结论", "方案具备实现可行性，但仍需补充实验验证"],
                ],
            ),
        ),
        (
            "benchmark-distractor.docx",
            _build_docx(
                paragraphs=[
                    "本资料讨论仓储搬运机器人系统。",
                    "重点是路径规划、机械臂抓取与分拣效率。",
                ],
                table_rows=[
                    ["题目", "仓储搬运机器人系统设计"],
                ],
            ),
        ),
    ]

    for filename, content in documents:
        response = client.post(
            f"/api/v1/projects/{project['id']}/sources/files",
            files=[
                (
                    "files",
                    (
                        filename,
                        content,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                )
            ],
        )
        if response.status_code != 201:
            raise RuntimeError(f"Failed to upload benchmark fixture {filename}: {response.text}")

    from app.repositories.search_repository import SearchRepository
    from app.repositories.source_repository import SourceRepository

    source_repository = SourceRepository()
    search_repository = SearchRepository()
    sources = source_repository.list_sources(project["id"])
    chunks = search_repository.get_latest_chunks_for_project(project["id"])
    return {
        "project": project,
        "sources": [source.to_summary() for source in sources],
        "chunks": chunks,
    }


def _build_resolved_cases(fixture: dict) -> list[ResolvedBenchmarkCase]:
    source_id_by_title = {source["title"]: source["id"] for source in fixture["sources"]}
    cases: list[ResolvedBenchmarkCase] = []
    for spec in _build_case_specs():
        chunk_ids = set()
        for matcher in spec.relevant_chunk_matchers:
            for chunk in fixture["chunks"]:
                if matcher.source_title and chunk["source_title"] != matcher.source_title:
                    continue
                if matcher.section_type and chunk.get("section_type") != matcher.section_type:
                    continue
                if matcher.field_label and chunk.get("field_label") != matcher.field_label:
                    continue
                if matcher.heading_path and chunk.get("heading_path") != matcher.heading_path:
                    continue
                if matcher.text_contains and matcher.text_contains not in (chunk.get("normalized_text") or ""):
                    continue
                chunk_ids.add(chunk["chunk_id"])

        source_ids = {source_id_by_title[title] for title in spec.relevant_source_titles if title in source_id_by_title}
        for chunk in fixture["chunks"]:
            if chunk["chunk_id"] in chunk_ids and chunk.get("source_id"):
                source_ids.add(chunk["source_id"])

        cases.append(
            ResolvedBenchmarkCase(
                case_id=spec.case_id,
                query=spec.query,
                query_type=spec.query_type,
                history=spec.history,
                relevant_chunk_ids=tuple(sorted(chunk_ids)),
                relevant_source_ids=tuple(sorted(source_ids)),
                notes=spec.notes,
            )
        )
    return cases


def _build_case_specs() -> list[BenchmarkCaseSpec]:
    specs: list[BenchmarkCaseSpec] = []
    direct_targets = [
        ("topic", ChunkMatcher(source_title="benchmark-overview.docx", section_type="field", field_label="题目")),
        ("project_name", ChunkMatcher(source_title="benchmark-overview.docx", section_type="field", field_label="项目名称")),
        ("innovation", ChunkMatcher(source_title="benchmark-overview.docx", section_type="field", field_label="创新点")),
        ("conclusion", ChunkMatcher(source_title="benchmark-feasibility.docx", section_type="field", field_label="结论")),
    ]
    direct_queries = {
        "topic": ["我的题目是什么？", "这份项目的题目叫什么？", "开题报告里的题目是什么？"],
        "project_name": ["项目名称是什么？", "这个系统叫什么名字？", "项目正式名称是什么？"],
        "innovation": ["创新点是什么？", "这份方案的创新点有哪些？", "项目的创新点写了什么？"],
        "conclusion": ["结论是什么？", "这份方案可行吗？", "结论部分怎么说？"],
    }
    for target, matcher in direct_targets:
        for index, query in enumerate(direct_queries[target], start=1):
            specs.append(BenchmarkCaseSpec(case_id=f"direct_{target}_{index}", query=query, query_type="direct_field", relevant_chunk_matchers=(matcher,)))

    follow_up_cases = [
        ("follow_up_topic_1", "现在你知道了吗？", [{"role": "user", "content_md": "我的题目是什么？"}], ChunkMatcher(source_title="benchmark-overview.docx", section_type="field", field_label="题目")),
        ("follow_up_topic_2", "那这份开题报告的题目呢？", [{"role": "user", "content_md": "你刚才说的是哪个项目？"}], ChunkMatcher(source_title="benchmark-overview.docx", section_type="field", field_label="题目")),
        ("follow_up_name_1", "那项目名称呢？", [{"role": "user", "content_md": "这份项目的题目叫什么？"}], ChunkMatcher(source_title="benchmark-overview.docx", section_type="field", field_label="项目名称")),
        ("follow_up_name_2", "现在你知道这个系统叫什么吗？", [{"role": "user", "content_md": "这个项目正式名称是什么？"}], ChunkMatcher(source_title="benchmark-overview.docx", section_type="field", field_label="项目名称")),
        ("follow_up_innovation_1", "那它的创新点呢？", [{"role": "user", "content_md": "项目名称是什么？"}], ChunkMatcher(source_title="benchmark-overview.docx", section_type="field", field_label="创新点")),
        ("follow_up_innovation_2", "现在能说出创新点了吗？", [{"role": "user", "content_md": "这份方案的创新点有哪些？"}], ChunkMatcher(source_title="benchmark-overview.docx", section_type="field", field_label="创新点")),
        ("follow_up_conclusion_1", "那结论怎么说？", [{"role": "user", "content_md": "这份方案可行吗？"}], ChunkMatcher(source_title="benchmark-feasibility.docx", section_type="field", field_label="结论")),
        ("follow_up_conclusion_2", "现在你知道可行性结论了吗？", [{"role": "user", "content_md": "结论是什么？"}], ChunkMatcher(source_title="benchmark-feasibility.docx", section_type="field", field_label="结论")),
    ]
    for case_id, query, history, matcher in follow_up_cases:
        specs.append(BenchmarkCaseSpec(case_id=case_id, query=query, query_type="natural_follow_up", history=history, relevant_chunk_matchers=(matcher,)))

    chapter_queries = [
        ("chapter_modules_1", "研究内容包括哪些模块？", ChunkMatcher(source_title="benchmark-overview.docx", text_contains="采集模块、控制模块、显示模块与报警模块")),
        ("chapter_modules_2", "实施内容里提到了哪些模块？", ChunkMatcher(source_title="benchmark-overview.docx", text_contains="采集模块、控制模块、显示模块与报警模块")),
        ("chapter_modules_3", "这套系统包含哪些功能模块？", ChunkMatcher(source_title="benchmark-overview.docx", text_contains="采集模块、控制模块、显示模块与报警模块")),
        ("chapter_control_1", "控制模块默认联动什么？", ChunkMatcher(source_title="benchmark-overview.docx", text_contains="联动风扇转速")),
        ("chapter_control_2", "控制模块会根据什么做联动？", ChunkMatcher(source_title="benchmark-overview.docx", text_contains="空气质量指标联动风扇转速")),
        ("chapter_control_3", "控制模块负责什么动作？", ChunkMatcher(source_title="benchmark-overview.docx", text_contains="联动风扇转速")),
        ("chapter_opt_1", "优化建议里主要建议了什么？", ChunkMatcher(source_title="benchmark-optimization.docx", text_contains="继续优化多传感器融合")),
        ("chapter_opt_2", "报告建议补充什么验证？", ChunkMatcher(source_title="benchmark-optimization.docx", text_contains="补充实验验证")),
        ("chapter_opt_3", "还建议减少什么问题？", ChunkMatcher(source_title="benchmark-optimization.docx", text_contains="降低误报率")),
        ("chapter_opt_4", "报告建议增加哪些测试？", ChunkMatcher(source_title="benchmark-optimization.docx", text_contains="长期稳定性测试")),
    ]
    for case_id, query, matcher in chapter_queries:
        specs.append(BenchmarkCaseSpec(case_id=case_id, query=query, query_type="chapter_theme", relevant_chunk_matchers=(matcher,)))

    complex_queries = [
        ("complex_1", "请总结这套方案的核心内容和可行性。", ("benchmark-overview.docx", "benchmark-feasibility.docx")),
        ("complex_2", "请概括项目主题、实现内容和结论。", ("benchmark-overview.docx", "benchmark-feasibility.docx")),
        ("complex_3", "请汇总研究内容和优化建议。", ("benchmark-overview.docx", "benchmark-optimization.docx")),
        ("complex_4", "请比较当前实现内容和后续优化方向。", ("benchmark-overview.docx", "benchmark-optimization.docx")),
        ("complex_5", "如果要写摘要，应该覆盖哪些重点？", ("benchmark-overview.docx", "benchmark-optimization.docx", "benchmark-feasibility.docx")),
        ("complex_6", "请总结这个项目的创新点、结论和不足。", ("benchmark-overview.docx", "benchmark-optimization.docx", "benchmark-feasibility.docx")),
        ("complex_7", "请说明方案为什么可行，以及还缺什么。", ("benchmark-feasibility.docx", "benchmark-optimization.docx")),
        ("complex_8", "请给我一个包含研究内容和优化建议的高层总结。", ("benchmark-overview.docx", "benchmark-optimization.docx")),
    ]
    for case_id, query, source_titles in complex_queries:
        specs.append(BenchmarkCaseSpec(case_id=case_id, query=query, query_type="complex_summary", relevant_source_titles=source_titles))

    unrelated_queries = ["今天北京天气如何？", "比特币今天多少钱？", "NBA 总决赛什么时候开始？", "请推荐一款咖啡机。", "如何办理美国签证？", "OpenAI 最新模型是什么？", "仓储机器人路径规划怎么做？", "怎么写 Python 爬虫？", "请帮我总结苹果发布会。", "如何学习高等数学？"]
    for index, query in enumerate(unrelated_queries, start=1):
        specs.append(BenchmarkCaseSpec(case_id=f"unrelated_{index}", query=query, query_type="unrelated"))
    return specs


def _build_docx(*, paragraphs: list[str], table_rows: list[list[str]]) -> bytes:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for col_index, value in enumerate(row):
                table.cell(row_index, col_index).text = value
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _copy_settings(settings: Settings, **updates) -> Settings:
    values = settings.model_dump() if hasattr(settings, "model_dump") else settings.dict()
    values.update(updates)
    return Settings(**values)


def _ordered_unique(items) -> list[str]:
    seen: set[str] = set()
    ordered: list[str] = []
    for item in items:
        if item in seen:
            continue
        seen.add(item)
        ordered.append(item)
    return ordered


def _count_by_query_type(cases: list[ResolvedBenchmarkCase]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for case in cases:
        counts[case.query_type] = counts.get(case.query_type, 0) + 1
    return counts


def _mean(values) -> float | None:
    filtered = [float(value) for value in values if value is not None]
    if not filtered:
        return None
    return round(sum(filtered) / len(filtered), 6)


def _percentile(values: list[float], ratio: float) -> float | None:
    if not values:
        return None
    ranked = sorted(values)
    index = max(0, min(len(ranked) - 1, math.ceil(len(ranked) * ratio) - 1))
    return round(float(ranked[index]), 6)
