from io import BytesIO
from pathlib import Path
import re

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from fastapi import HTTPException, UploadFile
from pypdf import PdfReader

from app.repositories.session_repository import SessionRepository
from app.repositories.source_repository import SourceRepository
from app.services.vector_store import VectorStore
from app.services.web_research_service import WebResearchService


class SourceService:
    def __init__(self) -> None:
        self.repository = SourceRepository()
        self.sessions = SessionRepository()
        self.vector_store = VectorStore()
        self.web_research = WebResearchService()

    def list_sources(self, project_id: str, *, include_archived: bool = False) -> list[dict]:
        if not self.repository.project_exists(project_id):
            raise HTTPException(status_code=404, detail=f"Project not found: {project_id}")
        return [record.to_summary() for record in self.repository.list_sources(project_id, include_archived=include_archived)]

    def list_all_sources(self, *, project_id: str | None = None, include_archived: bool = False) -> list[dict]:
        return [record.to_summary() for record in self.repository.list_all_sources(project_id=project_id, include_archived=include_archived)]

    def create_web_source(self, project_id: str, url: str, *, session_id: str | None = None) -> dict:
        if not self.repository.project_exists(project_id):
            raise HTTPException(status_code=404, detail=f"Project not found: {project_id}")

        try:
            fetched = self.web_research.fetch(url=url)
            canonical_uri = self.web_research.normalize_url(fetched["canonical_uri"])
            existing = self.repository.find_active_source_by_uri(project_id=project_id, canonical_uri=canonical_uri)
            if existing is not None:
                summary = existing.to_summary()
                if session_id is not None:
                    self._append_source_update_message(session_id, summary, "网页资料已在知识库中")
                return summary

            source = self.repository.create_web_source(project_id, canonical_uri)
            title = fetched["title"]
            text = fetched["text"]
            chunks = self._finalize_chunks(self._build_plain_text_chunks(text))
            self._complete_ingestion(
                source_id=source.id,
                project_id=project_id,
                title=title,
                text=text,
                chunks=chunks,
                reason="source_ingested",
            )
        except Exception as exc:
            raise HTTPException(status_code=422, detail=str(exc)) from exc

        refreshed = self.repository.get_source(source.id)
        if refreshed is None:
            raise HTTPException(status_code=500, detail="Source creation failed unexpectedly.")

        if session_id is not None:
            self._append_source_update_message(session_id, refreshed.to_summary(), "已添加网页资料")

        return refreshed.to_summary()

    async def create_file_sources(self, project_id: str, files: list[UploadFile], *, session_id: str | None = None) -> list[dict]:
        if not self.repository.project_exists(project_id):
            raise HTTPException(status_code=404, detail=f"Project not found: {project_id}")

        created_sources: list[dict] = []
        for file in files:
            filename = file.filename or "uploaded-file"
            suffix = Path(filename).suffix.lower()
            if suffix not in {".pdf", ".docx"}:
                raise HTTPException(status_code=422, detail=f"Unsupported file type: {filename}")

            source_type = "file_pdf" if suffix == ".pdf" else "file_docx"
            source = self.repository.create_file_source(
                project_id=project_id,
                source_type=source_type,
                title=filename,
                canonical_uri=f"file://{filename}",
                original_filename=filename,
                mime_type=file.content_type or "application/octet-stream",
            )

            try:
                content = await file.read()
                if source_type == "file_pdf":
                    text, chunks = self._extract_pdf_content(content)
                else:
                    text, chunks = self._extract_docx_content(content)
                self._complete_ingestion(
                    source_id=source.id,
                    project_id=project_id,
                    title=filename,
                    text=text,
                    chunks=chunks,
                    reason="source_ingested",
                )
            except Exception as exc:
                self.repository.finalize_source_failure(
                    source_id=source.id,
                    error_code="ingestion_failed",
                    error_message=str(exc),
                )

            refreshed = self.repository.get_source(source.id)
            if refreshed is not None:
                summary = refreshed.to_summary()
                created_sources.append(summary)
                if session_id is not None:
                    self._append_source_update_message(session_id, summary, "已添加文件资料")

        return created_sources

    def get_source_preview(self, source_id: str) -> dict:
        source = self.repository.get_source(source_id)
        if source is None or source.deleted_at is not None:
            raise HTTPException(status_code=404, detail=f"Source not found: {source_id}")

        preview_chunks = self.repository.get_source_preview_chunks(source_id)
        item = source.to_summary()
        item["preview_chunks"] = [chunk.to_summary() for chunk in preview_chunks]
        return item

    def refresh_source(self, source_id: str) -> dict:
        source = self.repository.get_source(source_id)
        if source is None or source.deleted_at is not None:
            raise HTTPException(status_code=404, detail=f"Source not found: {source_id}")
        if source.source_type != "web_page":
            raise HTTPException(status_code=422, detail="Only web sources can be refreshed.")

        self.repository.mark_source_processing(source_id)
        try:
            fetched = self.web_research.fetch(url=source.canonical_uri)
            title = fetched["title"]
            text = fetched["text"]
            chunks = self._finalize_chunks(self._build_plain_text_chunks(text))
            self._complete_ingestion(
                source_id=source.id,
                project_id=source.project_id,
                title=title,
                text=text,
                chunks=chunks,
                reason="source_refreshed",
            )
            if fetched["canonical_uri"] != source.canonical_uri:
                self.repository.update_web_source_url(source.id, fetched["canonical_uri"])
        except Exception as exc:
            self.repository.finalize_source_failure(
                source_id=source.id,
                error_code="refresh_failed",
                error_message=str(exc),
            )

        refreshed = self.repository.get_source(source.id)
        if refreshed is None:
            raise HTTPException(status_code=500, detail="Source refresh failed unexpectedly.")
        return refreshed.to_summary()

    def update_web_source(self, source_id: str, url: str) -> dict:
        try:
            updated = self.repository.update_web_source_url(source_id, self.web_research.normalize_url(url))
        except ValueError as exc:
            raise HTTPException(status_code=422, detail=str(exc)) from exc
        if updated is None:
            raise HTTPException(status_code=404, detail=f"Source not found: {source_id}")
        return updated.to_summary()

    def archive_source(self, source_id: str) -> dict:
        source = self.repository.archive_source(source_id)
        if source is None:
            raise HTTPException(status_code=404, detail=f"Source not found: {source_id}")
        self.vector_store.delete_source_points(source_id)
        return source.to_summary()

    def restore_source(self, source_id: str) -> dict:
        existing = self.repository.get_source(source_id)
        if existing is None or existing.deleted_at is not None:
            raise HTTPException(status_code=404, detail=f"Source not found: {source_id}")
        if existing.ingestion_status != "archived":
            raise HTTPException(status_code=422, detail="Only archived sources can be restored.")

        source = self.repository.restore_source(source_id)
        if source is None:
            raise HTTPException(status_code=404, detail=f"Source not found: {source_id}")
        self._sync_source_vectors(source_id)
        return source.to_summary()

    def delete_source(self, source_id: str) -> dict:
        existing = self.repository.get_source(source_id)
        if existing is None:
            raise HTTPException(status_code=404, detail=f"Source not found: {source_id}")
        source = self.repository.delete_source(source_id)
        if source is None:
            raise HTTPException(status_code=404, detail=f"Source not found: {source_id}")
        self.vector_store.delete_source_points(source_id)
        return source.to_summary()

    def _append_source_update_message(self, session_id: str, source: dict, prefix: str) -> None:
        detail = self.sessions.get_session_detail(session_id)
        if detail is None:
            return
        self.sessions.create_message(
            session_id=session_id,
            project_id=detail["project_id"],
            role="system",
            message_type="source_update",
            title=prefix,
            content_md=f"{prefix}：{source['title']}",
        )

    def _complete_ingestion(
        self,
        *,
        source_id: str,
        project_id: str,
        title: str,
        text: str,
        chunks: list[dict],
        reason: str,
    ) -> None:
        if not chunks:
            raise ValueError("No usable text could be extracted from the source.")

        quality_level = "normal" if len(chunks) >= 2 or len(text) >= 600 else "low"
        self.repository.complete_source_ingestion(
            source_id=source_id,
            project_id=project_id,
            title=title,
            quality_level=quality_level,
            chunks=chunks,
            reason=reason,
        )
        self._sync_source_vectors(source_id)

    def _extract_pdf_content(self, content: bytes) -> tuple[str, list[dict]]:
        reader = PdfReader(BytesIO(content))
        pages: list[str] = []
        blocks: list[dict] = []
        for page_number, page in enumerate(reader.pages, start=1):
            extracted = (page.extract_text() or "").strip()
            if extracted:
                pages.append(extracted)
                blocks.extend(self._extract_pdf_blocks(extracted, page_number=page_number))
        text = "\n\n".join(pages).strip()
        if not text:
            raise ValueError("The PDF did not produce readable text.")
        return text, self._finalize_chunks(self._build_structured_chunks(blocks))

    def _extract_docx_content(self, content: bytes) -> tuple[str, list[dict]]:
        document = Document(BytesIO(content))
        blocks = self._extract_docx_blocks(document)
        text = "\n\n".join(block["normalized_text"] for block in blocks if block["normalized_text"]).strip()
        if not text:
            raise ValueError("The DOCX did not produce readable text.")
        return text, self._finalize_chunks(self._build_structured_chunks(blocks))

    def _extract_docx_blocks(self, document) -> list[dict]:
        blocks: list[dict] = []
        heading_stack: list[tuple[int, str]] = []
        for item in self._iter_docx_block_items(document):
            if isinstance(item, Paragraph):
                paragraph_text = item.text.strip()
                if not paragraph_text:
                    continue
                heading_level = self._detect_docx_heading_level(item)
                if heading_level is not None:
                    heading_stack = [entry for entry in heading_stack if entry[0] < heading_level]
                    heading_stack.append((heading_level, paragraph_text))
                    blocks.append(
                        self._make_chunk_block(
                            text=paragraph_text,
                            section_label=paragraph_text,
                            section_type="heading",
                            heading_path=self._format_heading_path(heading_stack),
                        )
                    )
                    continue

                field_pair = self._split_field_pair(paragraph_text)
                if field_pair is not None:
                    label, value = field_pair
                    blocks.append(
                        self._make_chunk_block(
                            text=f"{label}: {value}",
                            section_label=label,
                            section_type="field",
                            heading_path=self._format_heading_path(heading_stack),
                            field_label=label,
                        )
                    )
                    continue

                blocks.append(
                    self._make_chunk_block(
                        text=paragraph_text,
                        section_label=self._default_section_label(heading_stack),
                        section_type="body",
                        heading_path=self._format_heading_path(heading_stack),
                    )
                )
                continue

            if isinstance(item, Table):
                table_blocks = self._extract_docx_table_blocks(item, heading_stack=heading_stack)
                blocks.extend(table_blocks)
        return blocks

    def _extract_docx_table_blocks(self, table: Table, *, heading_stack: list[tuple[int, str]]) -> list[dict]:
        blocks: list[dict] = []
        for row_index, row in enumerate(table.rows):
            cells: list[str] = []
            for cell in row.cells:
                cell_text = self._extract_docx_cell_text(cell)
                if cell_text:
                    cells.append(cell_text)
            if not cells:
                continue

            if len(cells) >= 2 and len(cells[0]) <= 24:
                label = cells[0]
                value = " ".join(cells[1:]).strip()
                blocks.append(
                    self._make_chunk_block(
                        text=f"{label}: {value}",
                        section_label=label,
                        section_type="field",
                        heading_path=self._format_heading_path(heading_stack),
                        field_label=label,
                        table_origin=f"table_row_{row_index + 1}",
                    )
                )
                continue

            row_text = " | ".join(cells)
            blocks.append(
                self._make_chunk_block(
                    text=row_text,
                    section_label=self._default_section_label(heading_stack),
                    section_type="table_row",
                    heading_path=self._format_heading_path(heading_stack),
                    table_origin=f"table_row_{row_index + 1}",
                )
            )
        return blocks

    def _extract_docx_cell_text(self, cell: _Cell) -> str:
        parts = [paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip()]
        for table in cell.tables:
            for row in table.rows:
                row_cells: list[str] = []
                for nested_cell in row.cells:
                    nested_text = " ".join(
                        paragraph.text.strip()
                        for paragraph in nested_cell.paragraphs
                        if paragraph.text.strip()
                    )
                    if nested_text:
                        row_cells.append(nested_text)
                if row_cells:
                    parts.append(" | ".join(row_cells))
        return " ".join(part for part in parts if part)

    def _fetch_web_content(self, url: str) -> tuple[str, str]:
        fetched = self.web_research.fetch(url=url)
        return fetched["title"], fetched["text"]

    def _build_plain_text_chunks(self, text: str, max_chars: int = 900) -> list[dict]:
        paragraphs = [paragraph.strip() for paragraph in text.splitlines() if paragraph.strip()]
        blocks = [
            self._make_chunk_block(
                text=paragraph,
                section_label="body",
                section_type="body",
            )
            for paragraph in paragraphs
        ]
        return self._build_structured_chunks(blocks, max_chars=max_chars)

    def _build_structured_chunks(self, blocks: list[dict], max_chars: int = 900) -> list[dict]:
        chunks: list[dict] = []
        current: dict | None = None

        def flush() -> None:
            nonlocal current
            if current is None:
                return
            normalized_text = current["normalized_text"].strip()
            if normalized_text:
                current["excerpt"] = normalized_text[:280]
                chunks.append(dict(current))
            current = None

        for block in blocks:
            text = block["normalized_text"].strip()
            if not text:
                continue
            if block["section_type"] in {"heading", "field", "table_row"}:
                flush()
                single = dict(block)
                single["normalized_text"] = text
                single["excerpt"] = text[:280]
                chunks.append(single)
                continue

            if current is None:
                current = dict(block)
                current["normalized_text"] = text
                continue

            same_bucket = (
                current["section_type"] == block["section_type"]
                and current.get("heading_path") == block.get("heading_path")
                and current.get("field_label") == block.get("field_label")
                and current.get("table_origin") == block.get("table_origin")
            )
            candidate = f"{current['normalized_text']}\n\n{text}".strip()
            if same_bucket and len(candidate) <= max_chars:
                current["normalized_text"] = candidate
                continue

            flush()
            current = dict(block)
            current["normalized_text"] = text

        flush()
        return chunks

    def _extract_pdf_blocks(self, text: str, *, page_number: int) -> list[dict]:
        blocks: list[dict] = []
        heading_stack: list[tuple[int, str]] = []
        for unit in self._split_pdf_units(text):
            line = " ".join(unit.split()).strip()
            if not line:
                continue

            heading_level = self._detect_pdf_heading_level(line)
            if heading_level is not None:
                heading_stack = [entry for entry in heading_stack if entry[0] < heading_level]
                heading_stack.append((heading_level, line))
                blocks.append(
                    self._make_chunk_block(
                        text=line,
                        section_label=line,
                        section_type="heading",
                        heading_path=self._format_heading_path(heading_stack),
                    )
                )
                continue

            if self._looks_like_heading_line(line):
                heading_stack = [(1, line)]
                blocks.append(
                    self._make_chunk_block(
                        text=line,
                        section_label=line,
                        section_type="heading",
                        heading_path=self._format_heading_path(heading_stack),
                    )
                )
                continue

            field_pair = self._split_field_pair(line)
            if field_pair is not None:
                label, value = field_pair
                blocks.append(
                    self._make_chunk_block(
                        text=f"{label}: {value}",
                        section_label=label,
                        section_type="field",
                        heading_path=self._format_heading_path(heading_stack),
                        field_label=label,
                    )
                )
                continue

            blocks.append(
                self._make_chunk_block(
                    text=line,
                    section_label=self._default_section_label(heading_stack) if heading_stack else f"page_{page_number}",
                    section_type="body",
                    heading_path=self._format_heading_path(heading_stack),
                )
            )
        return blocks

    def _iter_docx_block_items(self, parent):
        if hasattr(parent, "element") and hasattr(parent.element, "body"):
            parent_element = parent.element.body
        else:
            parent_element = parent._tc

        for child in parent_element.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def _detect_docx_heading_level(self, paragraph: Paragraph) -> int | None:
        style_name = (paragraph.style.name or "").lower() if paragraph.style is not None else ""
        if style_name.startswith("heading"):
            digits = "".join(character for character in style_name if character.isdigit())
            if digits:
                return int(digits)
            return 1

        text = paragraph.text.strip()
        if self._looks_like_heading_line(text):
            return 1
        return None

    def _detect_pdf_heading_level(self, text: str) -> int | None:
        numbered = re.match(r"^(?P<num>\d+(?:\.\d+){0,2})[\s、.．)\]）-]+", text)
        if numbered:
            return min(numbered.group("num").count(".") + 1, 3)

        chinese_numbered = re.match(r"^第?[一二三四五六七八九十百]+[章节部分篇][\s、.．)\]）-]*", text)
        if chinese_numbered:
            return 1
        return None

    def _looks_like_heading_line(self, text: str) -> bool:
        if len(text) > 32:
            return False
        if any(marker in text for marker in (":", "：", ".", "。", "?", "？", "!", "！", "|")):
            return False
        return True

    def _split_field_pair(self, text: str) -> tuple[str, str] | None:
        for separator in ("：", ":"):
            if separator not in text:
                continue
            label, value = text.split(separator, 1)
            label = label.strip()
            value = value.strip()
            if 0 < len(label) <= 24 and value:
                return label, value
        return None

    def _split_pdf_units(self, text: str) -> list[str]:
        paragraphs = [segment.strip() for segment in re.split(r"\n\s*\n+", text) if segment.strip()]
        if len(paragraphs) > 1:
            return paragraphs
        return [line.strip() for line in text.splitlines() if line.strip()]

    def _make_chunk_block(
        self,
        *,
        text: str,
        section_label: str,
        section_type: str,
        heading_path: str | None = None,
        field_label: str | None = None,
        table_origin: str | None = None,
        proposition_type: str | None = None,
    ) -> dict:
        normalized_text = " ".join(text.split()).strip()
        return {
            "section_label": section_label,
            "section_type": section_type,
            "heading_path": heading_path,
            "field_label": field_label,
            "table_origin": table_origin,
            "proposition_type": proposition_type or self._classify_proposition_type(
                text=normalized_text,
                field_label=field_label,
                heading_path=heading_path,
                section_type=section_type,
            ),
            "normalized_text": normalized_text,
            "excerpt": normalized_text[:280],
        }

    def _format_heading_path(self, heading_stack: list[tuple[int, str]]) -> str | None:
        if not heading_stack:
            return None
        return " > ".join(title for _, title in heading_stack)

    def _default_section_label(self, heading_stack: list[tuple[int, str]]) -> str:
        if not heading_stack:
            return "body"
        return heading_stack[-1][1]

    def _finalize_chunks(self, chunks: list[dict]) -> list[dict]:
        enriched = [dict(chunk) for chunk in chunks]
        enriched.extend(self._build_proposition_chunks(enriched))
        return enriched

    def _build_proposition_chunks(self, chunks: list[dict]) -> list[dict]:
        proposition_chunks: list[dict] = []
        seen: set[tuple[str, str | None, str | None, str]] = set()

        for chunk in chunks:
            if chunk.get("section_type") not in {"body", "field"}:
                continue
            sentences = self._extract_proposition_sentences(chunk["normalized_text"])
            for sentence in sentences:
                if sentence == chunk["normalized_text"]:
                    continue
                proposition_type = self._classify_proposition_type(
                    text=sentence,
                    field_label=chunk.get("field_label"),
                    heading_path=chunk.get("heading_path"),
                    section_type="proposition",
                )
                if proposition_type is None:
                    continue
                dedupe_key = (
                    sentence,
                    chunk.get("heading_path"),
                    chunk.get("field_label"),
                    proposition_type,
                )
                if dedupe_key in seen:
                    continue
                seen.add(dedupe_key)
                proposition_chunks.append(
                    {
                        "section_label": chunk["section_label"],
                        "section_type": "proposition",
                        "heading_path": chunk.get("heading_path"),
                        "field_label": chunk.get("field_label"),
                        "table_origin": chunk.get("table_origin"),
                        "proposition_type": proposition_type,
                        "normalized_text": sentence,
                        "excerpt": sentence[:280],
                    }
                )

        return proposition_chunks

    def _extract_proposition_sentences(self, text: str) -> list[str]:
        normalized = " ".join(text.split()).strip()
        if len(normalized) < 24:
            return []

        parts = re.split(r"(?<=[。！？；;.!?])\s+|(?<=[。！？；;.!?])", normalized)
        sentences: list[str] = []
        for part in parts:
            sentence = part.strip(" \t\r\n；;。.!?").strip()
            if len(sentence) < 12:
                continue
            if len(sentence) > 220:
                continue
            sentences.append(sentence)
        return sentences

    def _classify_proposition_type(
        self,
        *,
        text: str,
        field_label: str | None,
        heading_path: str | None,
        section_type: str,
    ) -> str | None:
        haystack = " ".join(part for part in (field_label, heading_path, text) if part).lower()

        if any(keyword in haystack for keyword in ("题目", "标题", "课题", "项目名称", "project name", "title")):
            return "identity"
        if any(keyword in haystack for keyword in ("建议", "优化", "改进", "推荐", "suggest", "recommend", "improve")):
            return "suggestion"
        if any(keyword in haystack for keyword in ("结论", "可行", "总结", "conclusion", "feasible")):
            return "conclusion"
        if any(keyword in haystack for keyword in ("创新", "novel", "innovation")):
            return "innovation"
        if any(keyword in haystack for keyword in ("预期成果", "成果", "deliverable", "outcome", "result")):
            return "outcome"
        if any(keyword in haystack for keyword in ("研究内容", "实施计划", "方法", "方案", "implementation", "method", "plan")):
            return "method"
        if section_type in {"body", "field", "proposition"} and len(text) >= 18:
            return "fact"
        return None

    def _sync_source_vectors(self, source_id: str) -> None:
        chunks = self.repository.get_latest_source_chunks_for_indexing(source_id)
        self.vector_store.upsert_source_chunks(chunks)
