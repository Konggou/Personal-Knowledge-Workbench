from __future__ import annotations

from dataclasses import asdict, dataclass, field
from io import BytesIO
from unittest.mock import patch

from docx import Document

from app.repositories.memory_repository import MemoryRepository
from app.repositories.session_repository import SessionRepository
from app.repositories.source_repository import SourceRepository
from app.services.agent_orchestrator_service import AgentOrchestratorService
from app.services.grounded_generation_service import GroundedGenerationService
from app.services.search_service import SearchService


@dataclass(frozen=True)
class RetrievalEvalCase:
    case_id: str
    query: str
    history: list[dict] | None = None
    limit: int = 3
    apply_rerank: bool = False
    expect_grounded_candidate: bool = True
    expect_second_pass: bool = False
    expect_min_hits: int = 1
    tags: tuple[str, ...] = field(default_factory=tuple)


@dataclass(frozen=True)
class MemorySeed:
    scope_type: str
    topic: str
    fact_text: str
    salience: float = 1.0


@dataclass(frozen=True)
class AgenticEvalCase:
    case_id: str
    query: str
    history: list[dict] | None = None
    research_mode: bool = False
    web_browsing: bool = False
    expect_grounded_candidate: bool = True
    expect_used_web: bool = False
    expect_readiness_action: str = "proceed"
    expect_second_pass: bool | None = None
    expect_min_evidence: int = 1
    expect_project_retry: bool = False
    expect_memory_notes: bool = False
    expect_primary_source_kind: str | None = "project_source"
    memory_seeds: tuple[MemorySeed, ...] = field(default_factory=tuple)
    mock_web_hits: tuple[dict, ...] = field(default_factory=tuple)
    force_empty_project_hits: bool = False
    force_project_retry: bool = False
    force_web_branch: bool = False
    tags: tuple[str, ...] = field(default_factory=tuple)


def build_retrieval_eval_fixture_docx() -> bytes:
    document = Document()
    document.add_heading("开题报告", level=1)
    document.add_paragraph("该系统面向室内空气质量检测与智能控制。")
    document.add_heading("研究内容", level=1)
    document.add_paragraph("实施计划包括采集模块、控制模块、显示模块与告警模块。")
    document.add_paragraph("控制模块负责根据空气质量指标联动风扇转速。")
    document.add_heading("优化建议", level=1)
    document.add_paragraph("可以进一步优化多传感器融合、降低误报率，并补充实验验证。")
    table = document.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "题目"
    table.cell(0, 1).text = "基于STM32的室内空气质量检测与智能控制系统设计"
    table.cell(1, 0).text = "项目名称"
    table.cell(1, 1).text = "室内空气质量检测与智能控制系统"
    table.cell(2, 0).text = "创新点"
    table.cell(2, 1).text = "多传感器融合与自动控制联动"
    table.cell(3, 0).text = "结论"
    table.cell(3, 1).text = "该方案具备实现可行性，但仍需补充实验验证。"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_retrieval_eval_cases() -> list[RetrievalEvalCase]:
    return [
        RetrievalEvalCase(
            case_id="direct_field_title",
            query="我的题目是什么？",
            expect_grounded_candidate=True,
            expect_second_pass=False,
            tags=("field", "title"),
        ),
        RetrievalEvalCase(
            case_id="natural_follow_up",
            query="现在你知道了吗？",
            history=[{"role": "user", "content_md": "我的开题报告题目是什么？"}],
            expect_grounded_candidate=True,
            expect_second_pass=True,
            tags=("follow_up", "contextual", "chinese"),
        ),
        RetrievalEvalCase(
            case_id="chapter_plan",
            query="研究内容里的实施计划是什么？",
            expect_grounded_candidate=True,
            expect_second_pass=False,
            tags=("chapter", "plan"),
        ),
        RetrievalEvalCase(
            case_id="chapter_suggestion",
            query="你觉得报告有哪些地方可以再优化？",
            expect_grounded_candidate=True,
            expect_second_pass=False,
            tags=("chapter", "suggestion", "complex"),
        ),
        RetrievalEvalCase(
            case_id="complex_grounded_non_research",
            query="请总结研究内容并说明为什么这个方案可行。",
            expect_grounded_candidate=True,
            expect_second_pass=False,
            apply_rerank=True,
            tags=("complex", "grounded"),
        ),
        RetrievalEvalCase(
            case_id="unrelated_weather",
            query="今天北京天气如何？",
            expect_grounded_candidate=False,
            expect_second_pass=True,
            expect_min_hits=0,
            tags=("unrelated", "weak_source_mode"),
        ),
    ]


def build_agentic_eval_cases() -> list[AgenticEvalCase]:
    return [
        AgenticEvalCase(
            case_id="project_only_grounding",
            query="我的题目是什么？",
            tags=("project_only", "field", "docx"),
        ),
        AgenticEvalCase(
            case_id="project_follow_up_second_pass",
            query="现在你知道了吗？",
            history=[{"role": "user", "content_md": "我的开题报告题目是什么？"}],
            expect_second_pass=True,
            tags=("follow_up", "contextual", "second_pass"),
        ),
        AgenticEvalCase(
            case_id="memory_assisted_follow_up",
            query="那这个方案可行吗？",
            memory_seeds=(
                MemorySeed(
                    scope_type="session",
                    topic="recent_answer",
                    fact_text="刚才已经确认过该方案具备实现可行性，但仍需补充实验验证。",
                ),
            ),
            expect_memory_notes=True,
            tags=("memory", "follow_up"),
        ),
        AgenticEvalCase(
            case_id="project_plus_web_supplement",
            query="请联网补充这套方案在行业里的 benchmark 结论。",
            web_browsing=True,
            expect_used_web=True,
            expect_primary_source_kind="external_web",
            mock_web_hits=(
                {
                    "source_kind": "external_web",
                    "source_title": "行业 Benchmark 摘要",
                    "source_type": "web_page",
                    "canonical_uri": "https://example.com/benchmark-summary",
                    "external_uri": "https://example.com/benchmark-summary",
                    "location_label": "网页补充 #1",
                    "excerpt": "外部资料提到类似系统的 benchmark 更关注稳定性和误报率控制。",
                    "normalized_text": "外部资料提到类似系统的 benchmark 更关注稳定性和误报率控制。",
                    "relevance_score": 3.9,
                    "section_type": "body",
                    "heading_path": None,
                    "field_label": None,
                    "table_origin": None,
                    "proposition_type": None,
                },
            ),
            force_empty_project_hits=True,
            tags=("web", "supplement"),
        ),
        AgenticEvalCase(
            case_id="precheck_retry_project_once",
            query="控制模块默认联动什么？",
            force_project_retry=True,
            expect_project_retry=True,
            tags=("precheck", "retry_project"),
        ),
        AgenticEvalCase(
            case_id="project_web_conflict_prefers_project",
            query="请联网补充并判断这个方案是否可行。",
            research_mode=True,
            web_browsing=True,
            expect_used_web=True,
            expect_primary_source_kind="project_source",
            force_web_branch=True,
            mock_web_hits=(
                {
                    "source_kind": "external_web",
                    "source_title": "外部博客观点",
                    "source_type": "web_page",
                    "canonical_uri": "https://example.com/conflict-note",
                    "external_uri": "https://example.com/conflict-note",
                    "location_label": "网页补充 #1",
                    "excerpt": "某外部博客认为这类方案尚不成熟，短期内难以落地。",
                    "normalized_text": "某外部博客认为这类方案尚不成熟，短期内难以落地。",
                    "relevance_score": 4.1,
                    "section_type": "body",
                    "heading_path": None,
                    "field_label": None,
                    "table_origin": None,
                    "proposition_type": None,
                },
            ),
            tags=("web", "conflict", "project_priority"),
        ),
        AgenticEvalCase(
            case_id="weak_source_mode_ready",
            query="今天北京天气如何？",
            expect_grounded_candidate=False,
            expect_readiness_action="insufficient",
            expect_min_evidence=0,
            expect_primary_source_kind=None,
            tags=("weak_source_mode", "fallback"),
        ),
    ]


def seed_retrieval_eval_project(client, *, project_name: str = "Retrieval Eval Project") -> dict:
    project_response = client.post(
        "/api/v1/projects",
        json={
            "name": project_name,
            "description": "Structured retrieval evaluation fixtures",
            "default_external_policy": "allow_external",
        },
    )
    if project_response.status_code != 201:
        raise RuntimeError(f"Failed to create eval project: {project_response.text}")
    project = project_response.json()["item"]

    upload_response = client.post(
        f"/api/v1/projects/{project['id']}/sources/files",
        files=[
            (
                "files",
                (
                    "retrieval-eval.docx",
                    build_retrieval_eval_fixture_docx(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            )
        ],
    )
    if upload_response.status_code != 201:
        raise RuntimeError(f"Failed to seed eval fixture source: {upload_response.text}")
    return project


def run_retrieval_eval(client, *, cases: list[RetrievalEvalCase] | None = None) -> dict:
    project = seed_retrieval_eval_project(client)
    service = SearchService()
    service.llm.is_configured = lambda: False
    grounded_generation = GroundedGenerationService(search_service=service)
    eval_cases = cases or build_retrieval_eval_cases()
    results: list[dict] = []

    for case in eval_cases:
        evidence, diagnostics = service.retrieve_project_evidence_with_diagnostics(
            project["id"],
            case.query,
            limit=case.limit,
            apply_rerank=case.apply_rerank,
            history=case.history,
        )
        packed_evidence = grounded_generation.retrieve_evidence(
            project_id=project["id"],
            query=case.query,
            research_mode=case.limit > 3,
            history=case.history,
        )
        delivery_diagnostics = grounded_generation.last_retrieval_diagnostics
        grounded_candidate = bool(diagnostics["final"]["grounded_candidate"])
        triggered_second_pass = bool(diagnostics["triggered_second_pass"])
        labels: list[str] = []

        if grounded_candidate != case.expect_grounded_candidate:
            labels.append("grounding_mismatch")
        if triggered_second_pass != case.expect_second_pass:
            labels.append("second_pass_mismatch")
        if grounded_candidate and len(evidence) < case.expect_min_hits:
            labels.append("insufficient_hits")
        if grounded_candidate and not packed_evidence:
            labels.append("delivery_empty")
        if not grounded_candidate and case.expect_grounded_candidate:
            labels.append("missed_grounding")
        if grounded_candidate and not case.expect_grounded_candidate:
            labels.append("false_positive_grounding")

        results.append(
            {
                **asdict(case),
                "grounded_candidate": grounded_candidate,
                "triggered_second_pass": triggered_second_pass,
                "hit_count": len(evidence),
                "packed_hit_count": len(packed_evidence),
                "source_count": diagnostics["final"]["source_count"],
                "top_score": diagnostics["first_pass"]["top_score"],
                "status": "passed" if not labels else "failed",
                "diagnostic_labels": labels,
                "diagnostics": diagnostics,
                "delivery_diagnostics": delivery_diagnostics,
                "evidence_titles": [item["source_title"] for item in evidence],
            }
        )

    return _finalize_report(project_id=project["id"], results=results)


def run_agentic_eval(client, *, cases: list[AgenticEvalCase] | None = None) -> dict:
    project = seed_retrieval_eval_project(client, project_name="Agentic Eval Project")
    repository = SourceRepository()
    session_repository = SessionRepository()
    memory_repository = MemoryRepository()
    source = repository.list_sources(project["id"])[0].to_summary()
    service = AgentOrchestratorService()
    service.llm.is_configured = lambda: False
    eval_cases = cases or build_agentic_eval_cases()
    results: list[dict] = []

    for case in eval_cases:
        session_response = client.post(f"/api/v1/projects/{project['id']}/sessions")
        if session_response.status_code != 201:
            raise RuntimeError(f"Failed to create eval session: {session_response.text}")
        session = session_response.json()["item"]
        _seed_case_memory(memory_repository, project_id=project["id"], session_id=session["id"], seeds=case.memory_seeds)

        history = _build_turn_history(case.history, case.query)
        patches: list = []
        if case.mock_web_hits:
            patches.append(
                patch.object(
                    service.web,
                    "build_external_evidence",
                    lambda **kwargs: _materialize_external_hits(project=project, hits=case.mock_web_hits),
                )
            )
        if case.force_empty_project_hits:
            patches.append(
                patch.object(
                    service.search,
                    "retrieve_project_evidence_with_diagnostics",
                    lambda *args, **kwargs: ([], _diagnostics_for([], grounded_candidate=False)),
                )
            )
        if case.force_project_retry:
            retrieval_sequence = _build_retry_retrieval_sequence(project=project, source=source)
            patches.append(
                patch.object(
                    service.search,
                    "retrieve_project_evidence_with_diagnostics",
                    _build_retry_retrieval_side_effect(retrieval_sequence),
                )
            )
            patches.append(
                patch.object(
                    service.llm,
                    "check_agent_answer_readiness",
                    _build_retry_readiness_side_effect(query=case.query),
                )
            )
        if case.force_web_branch:
            patches.append(patch.object(service, "_should_attempt_web", lambda _state: True))

        for active_patch in patches:
            active_patch.start()
        try:
            turn = service.orchestrate_turn(
                session_id=session["id"],
                project_id=project["id"],
                project_name=project["name"],
                query=case.query,
                history=history,
                research_mode=case.research_mode,
                web_browsing=case.web_browsing,
            )
        finally:
            while patches:
                patches.pop().stop()

        grounded_candidate = bool(turn.evidence_pack)
        primary_source_kind = _infer_source_kind(turn.evidence_pack[0]) if turn.evidence_pack else None
        diagnostics = turn.diagnostics or {}
        agentic_diagnostics = diagnostics.get("agentic", {})
        project_retry_count = int(agentic_diagnostics.get("project_retry_count", 0) or 0)
        memory_note_count = int(agentic_diagnostics.get("memory_note_count", 0) or 0)
        triggered_second_pass = bool(diagnostics.get("triggered_second_pass"))
        labels: list[str] = []

        if grounded_candidate != case.expect_grounded_candidate:
            labels.append("grounding_mismatch")
        if turn.used_web != case.expect_used_web:
            labels.append("web_usage_mismatch")
        if turn.readiness_action != case.expect_readiness_action:
            labels.append("readiness_action_mismatch")
        if case.expect_second_pass is not None and triggered_second_pass != case.expect_second_pass:
            labels.append("second_pass_mismatch")
        if grounded_candidate and len(turn.evidence_pack) < case.expect_min_evidence:
            labels.append("insufficient_evidence")
        if case.expect_project_retry and project_retry_count < 1:
            labels.append("project_retry_missing")
        if case.expect_memory_notes and memory_note_count < 1:
            labels.append("memory_lookup_missing")
        if primary_source_kind != case.expect_primary_source_kind:
            labels.append("primary_source_mismatch")

        results.append(
            {
                **asdict(case),
                "graph_profile": turn.graph_profile,
                "grounded_candidate": grounded_candidate,
                "used_web": turn.used_web,
                "readiness_action": turn.readiness_action,
                "triggered_second_pass": triggered_second_pass,
                "evidence_count": len(turn.evidence_pack),
                "memory_note_count": memory_note_count,
                "project_retry_count": project_retry_count,
                "primary_source_kind": primary_source_kind,
                "status": "passed" if not labels else "failed",
                "diagnostic_labels": labels,
                "diagnostics": diagnostics,
                "evidence_titles": [item["source_title"] for item in turn.evidence_pack],
                "query_trace": agentic_diagnostics.get("query_trace", []),
                "status_messages": turn.status_messages,
                "context_notes": turn.context_notes,
            }
        )

        session_repository.delete_session(session["id"])

    return _finalize_report(project_id=project["id"], results=results)


def run_v3_eval(client) -> dict:
    retrieval_report = run_retrieval_eval(client)
    agentic_report = run_agentic_eval(client)
    return {
        "suite": "v3",
        "passed_case_count": retrieval_report["passed_case_count"] + agentic_report["passed_case_count"],
        "failed_case_count": retrieval_report["failed_case_count"] + agentic_report["failed_case_count"],
        "retrieval": retrieval_report,
        "agentic": agentic_report,
    }


def _finalize_report(*, project_id: str, results: list[dict]) -> dict:
    passed_cases = sum(1 for item in results if item["status"] == "passed")
    return {
        "project_id": project_id,
        "case_count": len(results),
        "passed_case_count": passed_cases,
        "failed_case_count": len(results) - passed_cases,
        "results": results,
    }


def _build_turn_history(history: list[dict] | None, query: str) -> list[dict]:
    messages: list[dict] = []
    for index, item in enumerate(history or []):
        role = item["role"]
        messages.append(
            {
                "id": f"history-{index}",
                "role": "assistant" if role == "assistant" else "user",
                "message_type": "assistant_answer" if role == "assistant" else "user_prompt",
                "content_md": item["content_md"],
                "title": None,
                "sources": [],
            }
        )
    messages.append(
        {
            "id": "current-user-turn",
            "role": "user",
            "message_type": "user_prompt",
            "content_md": query,
            "title": None,
            "sources": [],
        }
    )
    return messages


def _seed_case_memory(
    repository: MemoryRepository,
    *,
    project_id: str,
    session_id: str,
    seeds: tuple[MemorySeed, ...],
) -> None:
    for seed in seeds:
        scope_id = session_id if seed.scope_type == "session" else project_id
        repository.upsert_entry(
            scope_type=seed.scope_type,
            scope_id=scope_id,
            topic=seed.topic,
            fact_text=seed.fact_text,
            salience=seed.salience,
            source_message_id=None,
        )


def _build_retry_retrieval_sequence(*, project: dict, source: dict) -> list[tuple[list[dict], dict]]:
    weak_hits = [
        {
            "project_id": project["id"],
            "project_name": project["name"],
            "chunk_id": None,
            "source_id": source["id"],
            "source_kind": "project_source",
            "source_title": source["title"],
            "source_type": source["source_type"],
            "canonical_uri": source["canonical_uri"],
            "external_uri": None,
            "location_label": "body #2",
            "excerpt": "控制模块会根据空气质量指标执行联动控制。",
            "normalized_text": "控制模块会根据空气质量指标执行联动控制。",
            "relevance_score": 2.6,
            "section_type": "body",
            "heading_path": "研究内容",
            "field_label": None,
            "table_origin": None,
            "proposition_type": None,
        }
    ]
    focused_hits = [
        {
            "project_id": project["id"],
            "project_name": project["name"],
            "chunk_id": None,
            "source_id": source["id"],
            "source_kind": "project_source",
            "source_title": source["title"],
            "source_type": source["source_type"],
            "canonical_uri": source["canonical_uri"],
            "external_uri": None,
            "location_label": "body #3",
            "excerpt": "控制模块负责根据空气质量指标联动风扇转速。",
            "normalized_text": "控制模块负责根据空气质量指标联动风扇转速。",
            "relevance_score": 4.2,
            "section_type": "proposition",
            "heading_path": "研究内容",
            "field_label": None,
            "table_origin": None,
            "proposition_type": "method",
        }
    ]
    return [
        (
            weak_hits,
            _diagnostics_for(
                weak_hits,
                grounded_candidate=True,
                top_score=2.6,
                term_coverage_ratio=0.2,
                input_candidate_count=4,
            ),
        ),
        (
            focused_hits,
            _diagnostics_for(
                focused_hits,
                grounded_candidate=True,
                top_score=4.2,
                term_coverage_ratio=0.8,
                input_candidate_count=2,
            ),
        ),
    ]


def _build_retry_retrieval_side_effect(sequence: list[tuple[list[dict], dict]]):
    state = {"index": 0}

    def _side_effect(*args, **kwargs):
        index = min(state["index"], len(sequence) - 1)
        state["index"] += 1
        return sequence[index]

    return _side_effect


def _build_retry_readiness_side_effect(*, query: str):
    state = {"count": 0}

    def _side_effect(**kwargs):
        state["count"] += 1
        if state["count"] == 1:
            return {
                "action": "retry_project",
                "reason": "project_evidence_not_focused",
                "focus": query,
            }
        return {
            "action": "proceed",
            "reason": "evidence_ready",
            "focus": "",
        }

    return _side_effect


def _materialize_external_hits(*, project: dict, hits: tuple[dict, ...]) -> list[dict]:
    materialized: list[dict] = []
    for index, item in enumerate(hits, start=1):
        materialized.append(
            {
                "project_id": project["id"],
                "project_name": project["name"],
                "chunk_id": None,
                "source_id": None,
                "source_kind": item["source_kind"],
                "source_title": item["source_title"],
                "source_type": item["source_type"],
                "canonical_uri": item["canonical_uri"],
                "external_uri": item["external_uri"],
                "location_label": item.get("location_label", f"网页补充 #{index}"),
                "excerpt": item["excerpt"],
                "normalized_text": item["normalized_text"],
                "relevance_score": item["relevance_score"],
                "section_type": item.get("section_type", "body"),
                "heading_path": item.get("heading_path"),
                "field_label": item.get("field_label"),
                "table_origin": item.get("table_origin"),
                "proposition_type": item.get("proposition_type"),
            }
        )
    return materialized


def _infer_source_kind(item: dict) -> str:
    if item.get("source_kind"):
        return str(item["source_kind"])
    return "project_source" if item.get("source_id") else "external_web"


def _diagnostics_for(
    evidence: list[dict],
    *,
    grounded_candidate: bool = True,
    top_score: float | None = None,
    term_coverage_ratio: float = 1.0,
    input_candidate_count: int | None = None,
) -> dict:
    inferred_top_score = top_score if top_score is not None else max(
        (float(item.get("relevance_score", 0.0)) for item in evidence),
        default=0.0,
    )
    return {
        "original_query": "eval",
        "context_clues": [],
        "first_pass": {
            "hit_count": len(evidence),
            "top_score": inferred_top_score,
            "title_hit_count": 1 if evidence else 0,
            "field_hit_count": sum(1 for item in evidence if item.get("section_type") == "field"),
            "term_coverage_ratio": term_coverage_ratio if evidence else 0.0,
            "is_low_confidence": not grounded_candidate,
        },
        "triggered_second_pass": False,
        "retry_steps": [],
        "final": {
            "hit_count": len(evidence),
            "source_count": len({item.get("source_id") or item.get("canonical_uri") for item in evidence}),
            "grounded_candidate": grounded_candidate and bool(evidence),
            "returned_hit_count": len(evidence),
            "selected_evidence_count": len(evidence),
        },
        "selection": {
            "input_candidate_count": input_candidate_count or len(evidence),
            "selected_candidate_count": len(evidence),
            "selector_applied": bool(input_candidate_count and input_candidate_count > len(evidence)),
            "items": [],
        },
        "compression": {
            "input_evidence_count": len(evidence),
            "compressed_evidence_count": len(evidence),
            "items": [],
        },
    }
